#!/usr/bin/env python3
"""生成多租戶開發計畫更新版報告"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    # Rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    return table


def main():
    doc = Document()

    # === 標題 ===
    title = doc.add_heading('小膳BOT 五公司多租戶開發計畫', level=0)
    doc.add_paragraph(f'更新日期：{datetime.now().strftime("%Y-%m-%d")}（民國 115 年 3 月 26 日）')
    doc.add_paragraph('版本：v3.0 — Phase 1 多租戶核心已完成')
    doc.add_paragraph('')

    # === Phase 1 完成狀態 ===
    add_heading(doc, '一、Phase 1 完成狀態', level=1)

    p = doc.add_paragraph()
    run = p.add_run('✅ Phase 1：多租戶核心 — 已完成')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 128, 0)

    add_table(doc,
        ['步驟', '工作項目', '狀態'],
        [
            ['1-1', '建立 companies 表 + 資料遷移腳本', '✅ 完成'],
            ['1-2', 'Webhook 多通道路由（Channel ID → company_id）', '✅ 完成'],
            ['1-3', 'LINE Service 多 Token 回覆', '✅ 完成'],
            ['1-4', 'GDrive Service 多路徑 + 自動建資料夾', '✅ 完成'],
            ['1-5', '所有查詢加 company_id 篩選', '✅ 完成'],
            ['1-6', 'Rich Menu 部署到五個帳號', '⏳ 待憑證設定後部署'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph('Phase 2（智慧分類）和 Phase 3（跨公司分析）待 Phase 1 穩定運行後開發。')

    # === 五家公司資料 ===
    add_heading(doc, '二、五家公司資料（已建入系統）', level=1)

    add_table(doc,
        ['#', 'LINE 帳號', '公司全名', '統一編號', '負責人', 'GDrive 資料夾', '憑證狀態'],
        [
            ['1', '福利社AI內帳', '升鼎商行', '81410187', '尤聖樺', '福利社/', '✅ 已設定'],
            ['2', '王凱食品AI內帳', '王凱食品有限公司', '90438334', '李佳芸', '王凱/', '❌ 待設定'],
            ['3', '台達2廠AI內帳', '升鼎商行', '81410187', '尤聖樺', '台達2廠/', '❌ 待設定'],
            ['4', '富燚AI內帳', '富燚商行', '00281384', '陳建男', '富燚/', '❌ 待設定'],
            ['5', '台達1廠AI內帳', '升鼎商行', '81410187', '尤聖樺', '台達1廠/', '❌ 待設定'],
        ]
    )

    # === GDrive 資料夾結構 ===
    add_heading(doc, '三、GDrive 資料夾結構（已建立）', level=1)

    doc.add_paragraph('五家公司的資料夾已在 Google 雲端建立完成：')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.style = 'No Spacing'
    p.add_run('小魚資料/團膳公司資料/').bold = True
    for line in [
        '├── 福利社/          ← 現有資料已搬入',
        '│   ├── 2025/',
        '│   ├── 2026/',
        '│   │   └── 03月/',
        '│   │       ├── 收據憑證/',
        '│   │       ├── 待確認/',
        '│   │       ├── 採購單據/',
        '│   │       ├── 月報表/',
        '│   │       └── ...（共 10 個子目錄）',
        '│   ├── 員工資料/',
        '│   ├── 教育訓練/',
        '│   └── 菜色圖片/',
        '├── 王凱/',
        '│   └── 2026/03月/（同上結構）',
        '├── 台達2廠/',
        '├── 富燚/',
        '└── 台達1廠/',
    ]:
        p = doc.add_paragraph(line)
        p.style = 'No Spacing'
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # === 小魚要做的事 ===
    add_heading(doc, '四、小魚需要完成的步驟', level=1)

    p = doc.add_paragraph()
    run = p.add_run('⚠️ 只差這一步，系統就能五家公司同時運作！')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)

    add_heading(doc, 'Step 1：登入 LINE Developers', level=2)
    doc.add_paragraph('1. 前往 https://developers.line.biz/')
    doc.add_paragraph('2. 用管理五個官方帳號的那個 LINE 帳號登入')

    add_heading(doc, 'Step 2：確認 Messaging API 已開啟', level=2)
    doc.add_paragraph('登入後應該能看到五個 Messaging API Channel。如果看不到：')
    doc.add_paragraph('1. 到 LINE Official Account Manager（https://manager.line.biz/）')
    doc.add_paragraph('2. 選擇要開啟的官方帳號')
    doc.add_paragraph('3. 設定 → Messaging API → 啟用')
    doc.add_paragraph('4. 選擇或建立 Provider')
    doc.add_paragraph('5. 完成後回到 LINE Developers Console 就能看到')

    add_heading(doc, 'Step 3：取得每個 Channel 的三組憑證', level=2)
    doc.add_paragraph('對「王凱」「台達2廠」「富燚」「台達1廠」四個 Channel 做以下操作：')
    doc.add_paragraph('')

    add_table(doc,
        ['憑證名稱', '在哪裡找', '長什麼樣'],
        [
            ['Channel ID', 'Basic settings 頁面 → Channel ID', '一串數字（如 2009160948）'],
            ['Channel Secret', '同頁面 → Channel secret', '32 字元英數字串'],
            ['Channel Access Token', 'Messaging API 分頁 → Channel access token → 點 Issue', '很長的英數字串（約 170 字元）'],
        ]
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('請將四家的憑證（共 12 個值）用以下格式傳給我：')
    run.bold = True

    doc.add_paragraph('')
    for company in ['王凱', '台達2廠', '富燚', '台達1廠']:
        p = doc.add_paragraph()
        p.style = 'No Spacing'
        p.add_run(f'【{company}】').bold = True
        for field in ['Channel ID：（貼上）', 'Channel Secret：（貼上）', 'Channel Access Token：（貼上）']:
            p = doc.add_paragraph(field)
            p.style = 'No Spacing'
        doc.add_paragraph('')

    add_heading(doc, 'Step 4：設定 Webhook URL', level=2)
    doc.add_paragraph('在每個 Channel 的 Messaging API 分頁：')
    doc.add_paragraph('1. Webhook URL → 填入：https://shanbot.kuangshin.tw/webhook')
    doc.add_paragraph('2. Use webhook → 開啟')
    doc.add_paragraph('3. 點 Verify 測試連線（應顯示 Success）')

    add_heading(doc, 'Step 5：關閉自動回覆', level=2)
    doc.add_paragraph('在 LINE Official Account Manager 每個帳號中：')
    doc.add_paragraph('1. 設定 → 回應設定')
    doc.add_paragraph('2. 自動回應訊息 → 關閉')
    doc.add_paragraph('3. Webhook → 開啟')

    # === 技術架構 ===
    add_heading(doc, '五、技術架構摘要', level=1)

    add_table(doc,
        ['項目', '說明'],
        [
            ['架構', '單實例多租戶（1 個 shanbot + 5 個 LINE Channel）'],
            ['資料庫', 'SQLite + company_id 欄位隔離（22 表 + 1 companies 表）'],
            ['GDrive', '每家公司獨立資料夾，標準 10 個月份子目錄'],
            ['Webhook', '統一入口 /webhook，由 Channel ID 自動路由'],
            ['版本', 'shanbot v3.0.0'],
            ['測試', '675 項自動測試全部通過'],
        ]
    )

    # === 驗收標準 ===
    add_heading(doc, '六、Phase 1 驗收標準', level=1)

    checks = [
        ('五家公司各自拍照上傳', '收據自動歸入對應公司的 GDrive 資料夾'),
        ('五家公司獨立記帳', '採購記錄帶 company_id，不會混帳'),
        ('五家公司獨立出報表', '月報表/稅務匯出按公司篩選'),
        ('心跳報告分列', '每日 16:00 報告按公司分列統計'),
        ('Rich Menu 各帳號一致', '六宮格選單部署到五個帳號'),
    ]
    for title, desc in checks:
        p = doc.add_paragraph()
        run = p.add_run(f'☐ {title}')
        run.bold = True
        p.add_run(f'：{desc}')

    # === 後續 ===
    add_heading(doc, '七、後續演進', level=1)

    doc.add_paragraph('Phase 1（當前）→ Phase 2 → Phase 3')
    doc.add_paragraph('多租戶核心     → 智慧分類    → 跨公司分析')
    doc.add_paragraph('5 LINE + 1 BOT → 統編自動判斷 → 合併報表')
    doc.add_paragraph('分開歸檔報表   → 供應商自動綁定 → 費用比較')

    # Save
    output_path = '/mnt/h/我的雲端硬碟/小魚資料/團膳公司資料/1150326_小膳BOT_五公司多租戶開發計畫_v3.0.docx'
    doc.save(output_path)
    print(f'✅ 報告已儲存：{output_path}')
    return output_path


if __name__ == '__main__':
    main()
