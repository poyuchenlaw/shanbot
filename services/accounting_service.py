"""小膳 Bot — 會計自動化服務（進階會計助理）

完整會計循環：
1. 進貨 → 複式分錄（借貸平衡）
2. 收入 → 分錄（含銷項稅額）
3. 薪資 → 分錄（含勞健保/勞退/代扣稅）
4. 固定資產折舊 → 分錄
5. 期末結帳（調整→結轉→開帳）
6. 財務報表（損益表/資產負債表/試算表）
7. 營業稅摘要（401申報用）
8. Excel 帳冊（含總帳/日記帳/財報）

遵循：台灣中小企業會計準則（EAS）+ 營業稅法 + 所得稅法
"""

import logging
import os
from datetime import datetime, timedelta

import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

import state_manager as sm

logger = logging.getLogger("shanbot.accounting")

# GDrive mount path
GDRIVE_BASE = os.environ.get("GDRIVE_MOUNT", "/mnt/h/我的雲端硬碟")
ACCOUNTING_DIR = os.path.join(GDRIVE_BASE, "小膳", "會計帳冊")

# 會計科目對照（進貨分類 → 會計科目）
CATEGORY_ACCOUNT_MAP = {
    "蔬菜": ("5110", "進貨—蔬菜類"),
    "肉類": ("5110", "進貨—肉類"),
    "水產": ("5110", "進貨—水產類"),
    "蛋豆": ("5110", "進貨—蛋豆類"),
    "乾貨": ("5110", "進貨—乾貨類"),
    "調味料": ("5110", "進貨—調味料"),
    "油品": ("5110", "進貨—油品類"),
    "米糧": ("5110", "進貨—米糧類"),
    "other": ("5110", "進貨—其他"),
    "其他": ("5110", "進貨—其他"),
}

# Excel 樣式
_HEADER_FONT = Font(name="微軟正黑體", size=11, bold=True, color="FFFFFF")
_HEADER_FILL = PatternFill(start_color="2E7D32", end_color="2E7D32", fill_type="solid")
_TITLE_FONT = Font(name="微軟正黑體", size=14, bold=True)
_NORMAL_FONT = Font(name="微軟正黑體", size=10)
_MONEY_FMT = '#,##0'
_DATE_FMT = 'YYYY-MM-DD'
_THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)
_CENTER = Alignment(horizontal="center", vertical="center")
_RIGHT = Alignment(horizontal="right", vertical="center")

# 藍色主題（用於費用/負債相關表格）
_BLUE_FILL = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
_LIGHT_GREEN = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
_LIGHT_RED = PatternFill(start_color="FFEBEE", end_color="FFEBEE", fill_type="solid")
_LIGHT_BLUE = PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid")
_LIGHT_YELLOW = PatternFill(start_color="FFF8E1", end_color="FFF8E1", fill_type="solid")


def _ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)


# =====================================================================
# 1. 進貨分錄（複式簿記核心）
# =====================================================================

def generate_journal_entries(staging_id: int) -> list[dict]:
    """根據確認後的 staging 自動生成複式分錄

    採購交易的分錄邏輯：
      借：進貨（按分類）      金額
      借：進項稅額            稅額（如可扣抵）
      貸：現金 / 應付帳款     總金額

    Returns: 生成的分錄列表
    """
    staging = sm.get_staging(staging_id)
    if not staging:
        logger.warning(f"Staging #{staging_id} not found")
        return []

    items = sm.get_purchase_items(staging_id)
    entry_date = staging.get("purchase_date") or datetime.now().strftime("%Y-%m-%d")
    year_month = staging.get("year_month") or entry_date[:7]
    supplier = staging.get("supplier_name") or "未知"
    total_amount = staging.get("total_amount") or 0
    tax_amount = staging.get("tax_amount") or 0
    subtotal = staging.get("subtotal") or (total_amount - tax_amount)
    deduction_code = staging.get("deduction_code") or "1"

    # 清除舊分錄（防重複）
    sm.delete_journal_entries_by_source("purchase", staging_id)

    entries = []

    # === 借方：進貨科目 ===
    total_debit = 0

    if items:
        # 按分類彙總金額
        cat_totals = {}
        for item in items:
            cat = item.get("category") or "other"
            amt = item.get("amount") or 0
            if cat not in cat_totals:
                cat_totals[cat] = 0
            cat_totals[cat] += amt

        items_total = sum(cat_totals.values())

        # 如果品項金額與 subtotal 有差異，按比例調整以確保平衡
        scale = (subtotal / items_total) if items_total > 0 else 1

        for cat, amt in cat_totals.items():
            if amt <= 0:
                continue
            scaled_amt = round(amt * scale, 0)
            code, name = CATEGORY_ACCOUNT_MAP.get(cat, ("5110", "進貨—其他"))
            eid = sm.add_journal_entry(
                entry_date=entry_date, year_month=year_month,
                source_type="purchase", source_id=staging_id,
                description=f"進貨-{supplier}（{cat}）",
                account_code=code, account_name=name,
                debit=scaled_amt, credit=0,
            )
            total_debit += scaled_amt
            entries.append({"id": eid, "side": "debit", "account": name, "amount": scaled_amt})
    else:
        # 沒有明細 → 整筆作為進貨
        eid = sm.add_journal_entry(
            entry_date=entry_date, year_month=year_month,
            source_type="purchase", source_id=staging_id,
            description=f"進貨-{supplier}",
            account_code="5110", account_name="進貨",
            debit=round(subtotal, 0), credit=0,
        )
        total_debit += round(subtotal, 0)
        entries.append({"id": eid, "side": "debit", "account": "進貨", "amount": subtotal})

    # === 借方：進項稅額（可扣抵時）===
    if tax_amount > 0 and deduction_code == "1":
        eid = sm.add_journal_entry(
            entry_date=entry_date, year_month=year_month,
            source_type="purchase", source_id=staging_id,
            description=f"進項稅額-{supplier}",
            account_code="1150", account_name="進項稅額",
            debit=round(tax_amount, 0), credit=0,
        )
        total_debit += round(tax_amount, 0)
        entries.append({"id": eid, "side": "debit", "account": "進項稅額", "amount": tax_amount})

    # === 貸方：現金（= 借方合計，確保借貸平衡）===
    credit_amount = total_debit

    eid = sm.add_journal_entry(
        entry_date=entry_date, year_month=year_month,
        source_type="purchase", source_id=staging_id,
        description=f"付款-{supplier}",
        account_code="1100", account_name="現金",
        debit=0, credit=credit_amount,
    )
    entries.append({"id": eid, "side": "credit", "account": "現金", "amount": credit_amount})

    logger.info(f"Generated {len(entries)} journal entries for staging #{staging_id}")
    return entries


def verify_balance(staging_id: int) -> dict:
    """驗證單筆交易的借貸平衡"""
    entries = sm.get_journal_entries_by_source("purchase", staging_id)
    total_debit = sum(e.get("debit", 0) for e in entries)
    total_credit = sum(e.get("credit", 0) for e in entries)
    balanced = abs(total_debit - total_credit) < 1  # 容許 $1 四捨五入差
    return {
        "staging_id": staging_id,
        "total_debit": total_debit,
        "total_credit": total_credit,
        "difference": total_debit - total_credit,
        "balanced": balanced,
        "entry_count": len(entries),
    }


# =====================================================================
# 2. 收入分錄（含銷項稅額）
# =====================================================================

def generate_income_journal_entries(income_id: int) -> list[dict]:
    """收入交易 → 複式分錄

    分錄邏輯（含稅收入）：
      借：現金/應收帳款    總金額
      貸：營業收入          未稅金額
      貸：銷項稅額          稅額（5%）

    如果是小規模營業人（免開發票），則不拆銷項稅額。
    """
    conn = sm._get_conn()
    row = conn.execute("SELECT * FROM income WHERE id=?", (income_id,)).fetchone()
    conn.close()
    if not row:
        return []

    income = dict(row)
    amount = income.get("amount", 0)
    if amount <= 0:
        return []

    year_month = income.get("year_month") or datetime.now().strftime("%Y-%m")
    entry_date = income.get("income_date") or datetime.now().strftime("%Y-%m-%d")
    desc = income.get("description") or "營業收入"
    source = income.get("source") or ""

    # 清除舊分錄
    sm.delete_journal_entries_by_source("income", income_id)

    entries = []

    # 判斷是否含稅（預設含稅 5%）
    tax_rate = 0.05
    subtotal = round(amount / (1 + tax_rate))
    tax_amount = amount - subtotal

    # 借方：現金（或應收帳款）
    debit_account = "1100" if source != "應收" else "1200"
    debit_name = "現金" if source != "應收" else "應收帳款"
    eid = sm.add_journal_entry(
        entry_date=entry_date, year_month=year_month,
        source_type="income", source_id=income_id,
        description=f"收款-{desc}",
        account_code=debit_account, account_name=debit_name,
        debit=amount, credit=0,
    )
    entries.append({"id": eid, "side": "debit", "account": debit_name, "amount": amount})

    # 貸方：營業收入（未稅）
    eid = sm.add_journal_entry(
        entry_date=entry_date, year_month=year_month,
        source_type="income", source_id=income_id,
        description=f"收入-{desc}",
        account_code="4100", account_name="營業收入",
        debit=0, credit=subtotal,
    )
    entries.append({"id": eid, "side": "credit", "account": "營業收入", "amount": subtotal})

    # 貸方：銷項稅額
    if tax_amount > 0:
        eid = sm.add_journal_entry(
            entry_date=entry_date, year_month=year_month,
            source_type="income", source_id=income_id,
            description=f"銷項稅額-{desc}",
            account_code="2150", account_name="銷項稅額",
            debit=0, credit=tax_amount,
        )
        entries.append({"id": eid, "side": "credit", "account": "銷項稅額", "amount": tax_amount})

    logger.info(f"Generated {len(entries)} income journal entries for income #{income_id}")
    return entries


# =====================================================================
# 3. 薪資分錄
# =====================================================================

def generate_payroll_journal_entries(year_month: str) -> list[dict]:
    """整月薪資 → 複式分錄

    分錄邏輯（每位員工）：
      借：薪資費用          底薪
      借：加班費            加班費
      借：伙食費            伙食津貼
      借：勞健保費用        雇主負擔勞健保
      借：勞退費用          雇主提繳 6%
        貸：現金              實發金額
        貸：代扣所得稅        所得稅
        貸：應付勞健保        員工自付勞健保
        貸：應付勞退          勞退自提
        貸：應付勞健保        雇主負擔勞健保（應繳）
        貸：應付勞退          雇主提繳勞退（應繳）
    """
    payrolls = sm.get_payroll_for_journal(year_month)
    if not payrolls:
        return []

    # 清除舊的薪資分錄
    sm.delete_journal_entries_by_source("payroll", 0)
    # 用 year_month 的最後一天作為分錄日期
    parts = year_month.split("-")
    y, m = int(parts[0]), int(parts[1])
    import calendar
    last_day = calendar.monthrange(y, m)[1]
    entry_date = f"{year_month}-{last_day}"

    all_entries = []

    # 彙總所有員工
    total_base = 0
    total_overtime = 0
    total_meal = 0
    total_bonus = 0
    total_labor_ins = 0
    total_health_ins = 0
    total_pension_self = 0
    total_income_tax = 0
    total_net = 0

    for p in payrolls:
        total_base += p.get("base_salary", 0) or 0
        total_overtime += p.get("overtime_pay", 0) or 0
        total_meal += p.get("meal_allowance", 0) or 0
        total_bonus += p.get("bonus", 0) or 0
        total_labor_ins += p.get("labor_insurance", 0) or 0
        total_health_ins += p.get("health_insurance", 0) or 0
        total_pension_self += p.get("pension_self", 0) or 0
        total_income_tax += p.get("income_tax", 0) or 0
        total_net += p.get("net_salary", 0) or 0

    total_gross = total_base + total_overtime + total_meal + total_bonus

    # 雇主負擔估算（勞保 70% + 健保 60%）
    employer_labor = round(total_labor_ins / 0.20 * 0.70) if total_labor_ins > 0 else 0
    employer_health = round(total_health_ins / 0.30 * 0.60) if total_health_ins > 0 else 0
    employer_pension = round(total_base * 0.06)  # 6% 雇主提繳

    # === 借方 ===
    debit_items = []
    if total_base > 0:
        debit_items.append(("6110", "薪資費用", total_base, "薪資費用"))
    if total_overtime > 0:
        debit_items.append(("6220", "加班費", total_overtime, "加班費"))
    if total_meal > 0:
        debit_items.append(("6230", "伙食費", total_meal, "伙食津貼"))
    if total_bonus > 0:
        debit_items.append(("6110", "薪資費用", total_bonus, "獎金"))
    if employer_labor + employer_health > 0:
        debit_items.append(("6200", "勞健保費用", employer_labor + employer_health, "雇主負擔勞健保"))
    if employer_pension > 0:
        debit_items.append(("6210", "勞退費用", employer_pension, "雇主提繳勞退"))

    total_debit = 0
    for code, name, amt, desc in debit_items:
        eid = sm.add_journal_entry(
            entry_date=entry_date, year_month=year_month,
            source_type="payroll", source_id=0,
            description=f"{year_month} {desc}",
            account_code=code, account_name=name,
            debit=amt, credit=0,
        )
        total_debit += amt
        all_entries.append({"id": eid, "side": "debit", "account": name, "amount": amt})

    # === 貸方 ===
    credit_items = []
    if total_net > 0:
        credit_items.append(("1100", "現金", total_net, "實發薪資"))
    if total_income_tax > 0:
        credit_items.append(("2310", "代扣所得稅", total_income_tax, "代扣所得稅"))
    if total_labor_ins + total_health_ins > 0:
        credit_items.append(("2210", "應付勞健保", total_labor_ins + total_health_ins, "員工自付勞健保"))
    if total_pension_self > 0:
        credit_items.append(("2220", "應付勞退", total_pension_self, "勞退自提"))
    # 雇主負擔部分也是貸方（應付）
    if employer_labor + employer_health > 0:
        credit_items.append(("2210", "應付勞健保", employer_labor + employer_health, "雇主負擔勞健保"))
    if employer_pension > 0:
        credit_items.append(("2220", "應付勞退", employer_pension, "雇主提繳勞退"))

    total_credit = 0
    for code, name, amt, desc in credit_items:
        eid = sm.add_journal_entry(
            entry_date=entry_date, year_month=year_month,
            source_type="payroll", source_id=0,
            description=f"{year_month} {desc}",
            account_code=code, account_name=name,
            debit=0, credit=amt,
        )
        total_credit += amt
        all_entries.append({"id": eid, "side": "credit", "account": name, "amount": amt})

    logger.info(f"Generated {len(all_entries)} payroll journal entries for {year_month}")
    return all_entries


# =====================================================================
# 4. 固定資產折舊
# =====================================================================

def generate_depreciation_entries(year_month: str) -> list[dict]:
    """計算當月折舊並生成分錄

    直線法：(成本 - 殘值) / 使用年限月數 = 月折舊額

    分錄：
      借：折舊費用    月折舊額
      貸：累計折舊    月折舊額
    """
    assets = sm.get_fixed_assets("active")
    if not assets:
        return []

    parts = year_month.split("-")
    y, m = int(parts[0]), int(parts[1])
    import calendar
    last_day = calendar.monthrange(y, m)[1]
    entry_date = f"{year_month}-{last_day}"

    # 清除舊折舊分錄
    sm.delete_journal_entries_by_source("depreciation", 0)

    entries = []
    total_depreciation = 0

    for asset in assets:
        cost = asset.get("cost", 0)
        salvage = asset.get("salvage_value", 0)
        months = asset.get("useful_life_months", 60) or 60
        if cost <= 0 or months <= 0:
            continue

        monthly_dep = round((cost - salvage) / months)
        if monthly_dep <= 0:
            continue

        # 檢查是否超過可折舊總額
        accumulated = asset.get("accumulated_depreciation", 0) or 0
        max_depreciable = cost - salvage
        if accumulated >= max_depreciable:
            continue
        if accumulated + monthly_dep > max_depreciable:
            monthly_dep = max_depreciable - accumulated

        total_depreciation += monthly_dep

        # 更新累計折舊
        sm.update_fixed_asset(asset["id"],
                              accumulated_depreciation=accumulated + monthly_dep)

    if total_depreciation > 0:
        # 借方：折舊費用
        eid = sm.add_journal_entry(
            entry_date=entry_date, year_month=year_month,
            source_type="depreciation", source_id=0,
            description=f"{year_month} 固定資產折舊",
            account_code="6160", account_name="折舊費用",
            debit=total_depreciation, credit=0,
        )
        entries.append({"id": eid, "side": "debit", "account": "折舊費用",
                        "amount": total_depreciation})

        # 貸方：累計折舊
        eid = sm.add_journal_entry(
            entry_date=entry_date, year_month=year_month,
            source_type="depreciation", source_id=0,
            description=f"{year_month} 累計折舊",
            account_code="1550", account_name="累計折舊",
            debit=0, credit=total_depreciation,
        )
        entries.append({"id": eid, "side": "credit", "account": "累計折舊",
                        "amount": total_depreciation})

    return entries


# =====================================================================
# 5. 期末結帳
# =====================================================================

def perform_period_end_closing(year_month: str) -> dict:
    """執行期末結帳

    步驟：
    1. 調整分錄（折舊等已在各模組處理）
    2. 結帳分錄：將收入/費用/成本結轉到「本期損益」
    3. 標記月份為已結帳

    結帳分錄邏輯：
      借：營業收入等（收入科目）    合計
      貸：本期損益                  合計
      借：本期損益                  合計
      貸：進貨/薪資/費用等          合計
    """
    # 檢查是否已結帳
    acct = sm.get_monthly_accounting(year_month)
    if acct and acct.get("is_closed"):
        return {"status": "already_closed", "message": f"{year_month} 已結帳"}

    # 先執行折舊
    generate_depreciation_entries(year_month)

    parts = year_month.split("-")
    y, m = int(parts[0]), int(parts[1])
    import calendar
    last_day = calendar.monthrange(y, m)[1]
    entry_date = f"{year_month}-{last_day}"

    # 取得試算表
    trial = sm.get_trial_balance(year_month)

    # 分類匯總
    total_revenue = 0  # 4xxx
    total_cost = 0     # 5xxx
    total_expense = 0  # 6xxx

    for t in trial:
        code = t.get("account_code", "")
        balance = t.get("balance", 0)  # debit - credit
        if code.startswith("4"):
            total_revenue += abs(balance)  # 收入 balance 通常為負（credit > debit）
        elif code.startswith("5"):
            total_cost += balance
        elif code.startswith("6"):
            total_expense += balance

    # 清除舊的結帳分錄
    sm.delete_journal_entries_by_source("closing", 0)

    entries = []

    # 結轉收入 → 本期損益（借：收入科目 / 貸：本期損益）
    if total_revenue > 0:
        for t in trial:
            code = t.get("account_code", "")
            if code.startswith("4") and abs(t["balance"]) > 0:
                sm.add_journal_entry(
                    entry_date=entry_date, year_month=year_month,
                    source_type="closing", source_id=0,
                    description=f"結轉 {t['account_name']}",
                    account_code=code, account_name=t["account_name"],
                    debit=abs(t["balance"]), credit=0,
                )
        sm.add_journal_entry(
            entry_date=entry_date, year_month=year_month,
            source_type="closing", source_id=0,
            description="結轉收入至本期損益",
            account_code="3300", account_name="本期損益",
            debit=0, credit=total_revenue,
        )
        entries.append({"type": "revenue_close", "amount": total_revenue})

    # 結轉成本+費用 → 本期損益（借：本期損益 / 貸：成本費用科目）
    total_costs_expenses = total_cost + total_expense
    if total_costs_expenses > 0:
        sm.add_journal_entry(
            entry_date=entry_date, year_month=year_month,
            source_type="closing", source_id=0,
            description="結轉成本費用至本期損益",
            account_code="3300", account_name="本期損益",
            debit=total_costs_expenses, credit=0,
        )
        for t in trial:
            code = t.get("account_code", "")
            if (code.startswith("5") or code.startswith("6")) and t["balance"] > 0:
                sm.add_journal_entry(
                    entry_date=entry_date, year_month=year_month,
                    source_type="closing", source_id=0,
                    description=f"結轉 {t['account_name']}",
                    account_code=code, account_name=t["account_name"],
                    debit=0, credit=t["balance"],
                )
        entries.append({"type": "expense_close", "amount": total_costs_expenses})

    net_income = total_revenue - total_costs_expenses

    # 更新月度會計總表
    sm.upsert_monthly_accounting(
        year_month,
        total_income=total_revenue,
        total_expense=total_costs_expenses,
        net_profit=net_income,
        is_closed=1,
    )

    return {
        "status": "closed",
        "year_month": year_month,
        "total_revenue": total_revenue,
        "total_cost": total_cost,
        "total_expense": total_expense,
        "net_income": net_income,
        "entries_count": len(entries),
    }


# =====================================================================
# 6. 財務報表
# =====================================================================

def generate_income_statement(year_month: str) -> dict:
    """損益表（Income Statement / P&L）

    營業收入
    - 營業成本（進貨等）
    = 營業毛利
    - 營業費用（薪資/租金/水電/折舊等）
    = 營業淨利
    """
    trial = sm.get_trial_balance(year_month)

    revenue = 0
    cost = 0
    expense = 0
    revenue_detail = []
    cost_detail = []
    expense_detail = []

    for t in trial:
        code = t.get("account_code", "")
        name = t.get("account_name", "")
        balance = t.get("balance", 0)

        if code.startswith("4"):
            amt = abs(balance)
            revenue += amt
            revenue_detail.append({"code": code, "name": name, "amount": amt})
        elif code.startswith("5"):
            cost += balance
            cost_detail.append({"code": code, "name": name, "amount": balance})
        elif code.startswith("6"):
            expense += balance
            expense_detail.append({"code": code, "name": name, "amount": balance})

    gross_profit = revenue - cost
    net_income = gross_profit - expense

    return {
        "year_month": year_month,
        "revenue": revenue,
        "revenue_detail": revenue_detail,
        "cost": cost,
        "cost_detail": cost_detail,
        "gross_profit": gross_profit,
        "expense": expense,
        "expense_detail": expense_detail,
        "net_income": net_income,
    }


def generate_balance_sheet(year_month: str) -> dict:
    """資產負債表（Balance Sheet）

    資產 = 負債 + 權益
    """
    trial = sm.get_trial_balance(year_month)

    assets = 0
    liabilities = 0
    equity = 0
    asset_detail = []
    liability_detail = []
    equity_detail = []

    for t in trial:
        code = t.get("account_code", "")
        name = t.get("account_name", "")
        balance = t.get("balance", 0)

        if code.startswith("1"):
            assets += balance
            asset_detail.append({"code": code, "name": name, "amount": balance})
        elif code.startswith("2"):
            amt = abs(balance)
            liabilities += amt
            liability_detail.append({"code": code, "name": name, "amount": amt})
        elif code.startswith("3"):
            amt = abs(balance)
            equity += amt
            equity_detail.append({"code": code, "name": name, "amount": amt})

    return {
        "year_month": year_month,
        "assets": assets,
        "asset_detail": asset_detail,
        "liabilities": liabilities,
        "liability_detail": liability_detail,
        "equity": equity,
        "equity_detail": equity_detail,
        "balanced": abs(assets - liabilities - equity) < 1,
    }


# =====================================================================
# 7. 營業稅摘要（401 申報用）
# =====================================================================

def get_vat_summary(tax_period: str) -> dict:
    """取得雙月營業稅摘要（401 申報用）

    tax_period 格式：'2026-03-04'（3月到4月）或直接 '2026-03'（自動判斷雙月）
    """
    if len(tax_period) == 7:
        # 自動判斷雙月期間
        parts = tax_period.split("-")
        y, m = int(parts[0]), int(parts[1])
        if m % 2 == 1:
            m1, m2 = m, m + 1
        else:
            m1, m2 = m - 1, m
        ym1 = f"{y}-{m1:02d}"
        ym2 = f"{y}-{m2:02d}"
    else:
        parts = tax_period.split("-")
        ym1 = f"{parts[0]}-{parts[1]}"
        ym2 = f"{parts[0]}-{parts[2]}"

    # 進項稅額（可扣抵）
    input_tax = 0
    input_count = 0
    for ym in [ym1, ym2]:
        stagings = sm.get_stagings_by_month(ym)
        for s in stagings:
            if s.get("deduction_code") == "1" and (s.get("tax_amount") or 0) > 0:
                input_tax += s.get("tax_amount", 0)
                input_count += 1

    # 銷項稅額
    output_tax = 0
    output_count = 0
    for ym in [ym1, ym2]:
        entries = sm.get_journal_entries(ym, source_type="income")
        for e in entries:
            if e.get("account_code") == "2150":
                output_tax += e.get("credit", 0)
                output_count += 1

    # 應納/溢付
    tax_payable = output_tax - input_tax

    return {
        "tax_period": f"{ym1} ~ {ym2}",
        "output_tax": output_tax,
        "output_count": output_count,
        "input_tax": input_tax,
        "input_count": input_count,
        "tax_payable": tax_payable,
        "status": "應繳" if tax_payable > 0 else "留抵",
    }


# =====================================================================
# 8. 日記帳 Excel 生成（完整帳冊）
# =====================================================================

def generate_accounting_excel(year_month: str) -> str | None:
    """生成月度會計 Excel（完整帳冊）

    Sheet 1: 進貨日記帳
    Sheet 2: 月度費用彙總
    Sheet 3: 試算表
    Sheet 4: 分錄明細
    Sheet 5: 收入明細
    Sheet 6: 損益表
    Sheet 7: 資產負債表
    Sheet 8: 總分類帳

    Returns: Excel 檔案路徑
    """
    out_dir = os.path.join(ACCOUNTING_DIR, year_month)
    _ensure_dir(out_dir)
    filepath = os.path.join(out_dir, f"{year_month}_會計帳冊.xlsx")

    stagings = sm.get_stagings_by_month(year_month)
    journal_entries = sm.get_journal_entries(year_month)
    trial_balance = sm.get_trial_balance(year_month)
    income_rows = sm.get_income_summary(year_month)

    wb = openpyxl.Workbook()

    # --- Sheet 1: 進貨日記帳 ---
    ws1 = wb.active
    ws1.title = "進貨日記帳"
    _write_purchase_journal(ws1, stagings, year_month)

    # --- Sheet 2: 月度費用彙總 ---
    ws2 = wb.create_sheet("月度費用彙總")
    _write_expense_summary(ws2, stagings, year_month)

    # --- Sheet 3: 試算表（借貸平衡驗證）---
    ws3 = wb.create_sheet("試算表")
    _write_trial_balance(ws3, trial_balance, year_month)

    # --- Sheet 4: 分錄明細 ---
    ws4 = wb.create_sheet("分錄明細")
    _write_journal_entries_sheet(ws4, journal_entries, year_month)

    # --- Sheet 5: 收入明細 ---
    if income_rows:
        ws5 = wb.create_sheet("收入明細")
        _write_income_sheet(ws5, income_rows, year_month)

    # --- Sheet 6: 損益表 ---
    ws6 = wb.create_sheet("損益表")
    pl_data = generate_income_statement(year_month)
    _write_income_statement_sheet(ws6, pl_data)

    # --- Sheet 7: 資產負債表 ---
    ws7 = wb.create_sheet("資產負債表")
    bs_data = generate_balance_sheet(year_month)
    _write_balance_sheet(ws7, bs_data)

    # --- Sheet 8: 總分類帳 ---
    ws8 = wb.create_sheet("總分類帳")
    ledger = sm.get_general_ledger(year_month)
    _write_general_ledger_sheet(ws8, ledger, year_month)

    wb.save(filepath)
    logger.info(f"Accounting Excel generated: {filepath}")

    # 更新月度會計總表
    _update_monthly_accounting(year_month, filepath, journal_entries, income_rows, stagings)

    return filepath


# === Excel 寫入輔助函數 ===

def _write_purchase_journal(ws, stagings: list, year_month: str):
    """Sheet 1: 進貨日記帳明細"""
    ws.merge_cells("A1:H1")
    ws["A1"] = f"進貨日記帳 — {year_month}"
    ws["A1"].font = _TITLE_FONT

    headers = ["日期", "供應商", "項目名稱", "數量", "單位", "單價", "總價", "備註"]
    widths = [12, 16, 20, 8, 6, 10, 12, 20]

    for col_idx, (header, width) in enumerate(zip(headers, widths), 1):
        cell = ws.cell(row=3, column=col_idx, value=header)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _CENTER
        cell.border = _THIN_BORDER
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    row = 4
    grand_total = 0

    for staging in stagings:
        items = sm.get_purchase_items(staging["id"])
        supplier = staging.get("supplier_name") or "未知"
        pdate = staging.get("purchase_date") or ""
        notes = staging.get("notes") or ""

        if items:
            for item in items:
                amt = item.get("amount") or 0
                ws.cell(row=row, column=1, value=pdate).font = _NORMAL_FONT
                ws.cell(row=row, column=2, value=supplier).font = _NORMAL_FONT
                ws.cell(row=row, column=3, value=item.get("item_name", "")).font = _NORMAL_FONT
                ws.cell(row=row, column=4, value=item.get("quantity", 0)).font = _NORMAL_FONT
                ws.cell(row=row, column=5, value=item.get("unit", "")).font = _NORMAL_FONT

                price_cell = ws.cell(row=row, column=6, value=item.get("unit_price", 0))
                price_cell.font = _NORMAL_FONT
                price_cell.number_format = _MONEY_FMT
                price_cell.alignment = _RIGHT

                amt_cell = ws.cell(row=row, column=7, value=amt)
                amt_cell.font = _NORMAL_FONT
                amt_cell.number_format = _MONEY_FMT
                amt_cell.alignment = _RIGHT

                ws.cell(row=row, column=8, value=notes).font = _NORMAL_FONT

                for c in range(1, 9):
                    ws.cell(row=row, column=c).border = _THIN_BORDER
                grand_total += amt
                row += 1
        else:
            amt = staging.get("total_amount") or 0
            ws.cell(row=row, column=1, value=pdate).font = _NORMAL_FONT
            ws.cell(row=row, column=2, value=supplier).font = _NORMAL_FONT
            ws.cell(row=row, column=3, value="（整筆進貨）").font = _NORMAL_FONT
            for c in [4, 5, 6]:
                ws.cell(row=row, column=c, value="").font = _NORMAL_FONT

            amt_cell = ws.cell(row=row, column=7, value=amt)
            amt_cell.font = _NORMAL_FONT
            amt_cell.number_format = _MONEY_FMT
            amt_cell.alignment = _RIGHT

            ws.cell(row=row, column=8, value=notes).font = _NORMAL_FONT
            for c in range(1, 9):
                ws.cell(row=row, column=c).border = _THIN_BORDER
            grand_total += amt
            row += 1

    row += 1
    ws.cell(row=row, column=6, value="合計").font = Font(name="微軟正黑體", size=11, bold=True)
    total_cell = ws.cell(row=row, column=7, value=grand_total)
    total_cell.font = Font(name="微軟正黑體", size=11, bold=True)
    total_cell.number_format = _MONEY_FMT
    total_cell.alignment = _RIGHT


def _write_expense_summary(ws, stagings: list, year_month: str):
    """Sheet 2: 月度費用彙總"""
    ws.merge_cells("A1:D1")
    ws["A1"] = f"月度費用彙總 — {year_month}"
    ws["A1"].font = _TITLE_FONT

    ws.cell(row=3, column=1, value="【按供應商彙總】").font = Font(
        name="微軟正黑體", size=11, bold=True)

    headers_s = ["供應商", "筆數", "金額", "稅額"]
    for col_idx, h in enumerate(headers_s, 1):
        cell = ws.cell(row=4, column=col_idx, value=h)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _CENTER
        cell.border = _THIN_BORDER
    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 8
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 12

    supplier_totals = {}
    for s in stagings:
        name = s.get("supplier_name") or "未知"
        if name not in supplier_totals:
            supplier_totals[name] = {"count": 0, "amount": 0, "tax": 0}
        supplier_totals[name]["count"] += 1
        supplier_totals[name]["amount"] += s.get("total_amount") or 0
        supplier_totals[name]["tax"] += s.get("tax_amount") or 0

    row = 5
    for name, data in sorted(supplier_totals.items(), key=lambda x: -x[1]["amount"]):
        ws.cell(row=row, column=1, value=name).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=data["count"]).font = _NORMAL_FONT
        c = ws.cell(row=row, column=3, value=data["amount"])
        c.font = _NORMAL_FONT
        c.number_format = _MONEY_FMT
        c.alignment = _RIGHT
        t = ws.cell(row=row, column=4, value=data["tax"])
        t.font = _NORMAL_FONT
        t.number_format = _MONEY_FMT
        t.alignment = _RIGHT
        for col in range(1, 5):
            ws.cell(row=row, column=col).border = _THIN_BORDER
        row += 1

    # 按分類彙總
    row += 2
    ws.cell(row=row, column=1, value="【按分類彙總】").font = Font(
        name="微軟正黑體", size=11, bold=True)
    row += 1

    headers_c = ["分類", "金額", "佔比"]
    for col_idx, h in enumerate(headers_c, 1):
        cell = ws.cell(row=row, column=col_idx, value=h)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _CENTER
        cell.border = _THIN_BORDER

    cat_totals = {}
    total_amount = 0
    for s in stagings:
        items = sm.get_purchase_items(s["id"])
        for item in items:
            cat = item.get("category") or "其他"
            amt = item.get("amount") or 0
            cat_totals[cat] = cat_totals.get(cat, 0) + amt
            total_amount += amt

    row += 1
    for cat, amt in sorted(cat_totals.items(), key=lambda x: -x[1]):
        pct = (amt / total_amount * 100) if total_amount > 0 else 0
        ws.cell(row=row, column=1, value=cat).font = _NORMAL_FONT
        c = ws.cell(row=row, column=2, value=amt)
        c.font = _NORMAL_FONT
        c.number_format = _MONEY_FMT
        c.alignment = _RIGHT
        p = ws.cell(row=row, column=3, value=f"{pct:.1f}%")
        p.font = _NORMAL_FONT
        p.alignment = _RIGHT
        for col in range(1, 4):
            ws.cell(row=row, column=col).border = _THIN_BORDER
        row += 1


def _write_trial_balance(ws, trial_balance: list, year_month: str):
    """Sheet 3: 試算表"""
    ws.merge_cells("A1:E1")
    ws["A1"] = f"試算表 — {year_month}"
    ws["A1"].font = _TITLE_FONT

    headers = ["科目代碼", "科目名稱", "借方合計", "貸方合計", "餘額"]
    widths = [12, 20, 14, 14, 14]

    for col_idx, (h, w) in enumerate(zip(headers, widths), 1):
        cell = ws.cell(row=3, column=col_idx, value=h)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _CENTER
        cell.border = _THIN_BORDER
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    row = 4
    sum_debit = 0
    sum_credit = 0

    for tb in trial_balance:
        ws.cell(row=row, column=1, value=tb["account_code"]).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=tb["account_name"]).font = _NORMAL_FONT

        d = ws.cell(row=row, column=3, value=tb["total_debit"])
        d.font = _NORMAL_FONT
        d.number_format = _MONEY_FMT
        d.alignment = _RIGHT

        c = ws.cell(row=row, column=4, value=tb["total_credit"])
        c.font = _NORMAL_FONT
        c.number_format = _MONEY_FMT
        c.alignment = _RIGHT

        b = ws.cell(row=row, column=5, value=tb["balance"])
        b.font = _NORMAL_FONT
        b.number_format = _MONEY_FMT
        b.alignment = _RIGHT

        for col in range(1, 6):
            ws.cell(row=row, column=col).border = _THIN_BORDER

        sum_debit += tb["total_debit"]
        sum_credit += tb["total_credit"]
        row += 1

    row += 1
    ws.cell(row=row, column=2, value="合計").font = Font(name="微軟正黑體", size=11, bold=True)
    sd = ws.cell(row=row, column=3, value=sum_debit)
    sd.font = Font(name="微軟正黑體", size=11, bold=True)
    sd.number_format = _MONEY_FMT
    sd.alignment = _RIGHT

    sc = ws.cell(row=row, column=4, value=sum_credit)
    sc.font = Font(name="微軟正黑體", size=11, bold=True)
    sc.number_format = _MONEY_FMT
    sc.alignment = _RIGHT

    diff = sum_debit - sum_credit
    diff_cell = ws.cell(row=row, column=5, value=diff)
    diff_cell.font = Font(name="微軟正黑體", size=11, bold=True,
                          color="FF0000" if abs(diff) > 0.5 else "008000")
    diff_cell.number_format = _MONEY_FMT
    diff_cell.alignment = _RIGHT

    row += 2
    balanced = abs(diff) < 1
    status = "✅ 借貸平衡" if balanced else "❌ 借貸不平衡！差額 ${:,.0f}".format(diff)
    ws.cell(row=row, column=1, value=status).font = Font(
        name="微軟正黑體", size=12, bold=True,
        color="008000" if balanced else "FF0000")


def _write_journal_entries_sheet(ws, entries: list, year_month: str):
    """Sheet 4: 分錄明細"""
    ws.merge_cells("A1:G1")
    ws["A1"] = f"分錄明細（日記帳）— {year_month}"
    ws["A1"].font = _TITLE_FONT

    headers = ["日期", "摘要", "科目代碼", "科目名稱", "借方", "貸方", "來源"]
    widths = [12, 24, 10, 16, 12, 12, 10]

    for col_idx, (h, w) in enumerate(zip(headers, widths), 1):
        cell = ws.cell(row=3, column=col_idx, value=h)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _CENTER
        cell.border = _THIN_BORDER
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    row = 4
    for e in entries:
        ws.cell(row=row, column=1, value=e.get("entry_date", "")).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=e.get("description", "")).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=e.get("account_code", "")).font = _NORMAL_FONT
        ws.cell(row=row, column=4, value=e.get("account_name", "")).font = _NORMAL_FONT

        d = ws.cell(row=row, column=5, value=e.get("debit", 0) or "")
        d.font = _NORMAL_FONT
        if e.get("debit"):
            d.number_format = _MONEY_FMT
        d.alignment = _RIGHT

        c = ws.cell(row=row, column=6, value=e.get("credit", 0) or "")
        c.font = _NORMAL_FONT
        if e.get("credit"):
            c.number_format = _MONEY_FMT
        c.alignment = _RIGHT

        source = f"{e.get('source_type', '')}#{e.get('source_id', '')}"
        ws.cell(row=row, column=7, value=source).font = _NORMAL_FONT

        for col in range(1, 8):
            ws.cell(row=row, column=col).border = _THIN_BORDER
        row += 1


def _write_income_sheet(ws, income_rows: list, year_month: str):
    """Sheet 5: 收入明細"""
    ws.merge_cells("A1:E1")
    ws["A1"] = f"收入明細 — {year_month}"
    ws["A1"].font = _TITLE_FONT

    headers = ["日期", "說明", "來源", "金額", "備註"]
    widths = [12, 24, 12, 14, 16]

    for col_idx, (h, w) in enumerate(zip(headers, widths), 1):
        cell = ws.cell(row=3, column=col_idx, value=h)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _CENTER
        cell.border = _THIN_BORDER
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    row = 4
    total = 0
    for inc in income_rows:
        ws.cell(row=row, column=1, value=inc.get("income_date", "")).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=inc.get("description", "")).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=inc.get("source", "")).font = _NORMAL_FONT
        a = ws.cell(row=row, column=4, value=inc.get("amount", 0))
        a.font = _NORMAL_FONT
        a.number_format = _MONEY_FMT
        a.alignment = _RIGHT
        ws.cell(row=row, column=5, value="").font = _NORMAL_FONT
        for col in range(1, 6):
            ws.cell(row=row, column=col).border = _THIN_BORDER
        total += inc.get("amount", 0)
        row += 1

    row += 1
    ws.cell(row=row, column=3, value="合計").font = Font(name="微軟正黑體", size=11, bold=True)
    t = ws.cell(row=row, column=4, value=total)
    t.font = Font(name="微軟正黑體", size=11, bold=True)
    t.number_format = _MONEY_FMT
    t.alignment = _RIGHT


def _write_income_statement_sheet(ws, pl_data: dict):
    """Sheet 6: 損益表"""
    ym = pl_data["year_month"]
    ws.merge_cells("A1:C1")
    ws["A1"] = f"損益表 — {ym}"
    ws["A1"].font = _TITLE_FONT

    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 16

    row = 3
    bold = Font(name="微軟正黑體", size=11, bold=True)

    # 營業收入
    ws.cell(row=row, column=1, value="一、營業收入").font = bold
    ws.cell(row=row, column=1).fill = _LIGHT_GREEN
    row += 1
    for d in pl_data["revenue_detail"]:
        ws.cell(row=row, column=1, value=f"  {d['name']}").font = _NORMAL_FONT
        c = ws.cell(row=row, column=2, value=d["amount"])
        c.font = _NORMAL_FONT
        c.number_format = _MONEY_FMT
        c.alignment = _RIGHT
        row += 1
    ws.cell(row=row, column=1, value="  營業收入合計").font = bold
    c = ws.cell(row=row, column=3, value=pl_data["revenue"])
    c.font = bold
    c.number_format = _MONEY_FMT
    c.alignment = _RIGHT
    row += 2

    # 營業成本
    ws.cell(row=row, column=1, value="二、營業成本").font = bold
    ws.cell(row=row, column=1).fill = _LIGHT_RED
    row += 1
    for d in pl_data["cost_detail"]:
        ws.cell(row=row, column=1, value=f"  {d['name']}").font = _NORMAL_FONT
        c = ws.cell(row=row, column=2, value=d["amount"])
        c.font = _NORMAL_FONT
        c.number_format = _MONEY_FMT
        c.alignment = _RIGHT
        row += 1
    ws.cell(row=row, column=1, value="  營業成本合計").font = bold
    c = ws.cell(row=row, column=3, value=pl_data["cost"])
    c.font = bold
    c.number_format = _MONEY_FMT
    c.alignment = _RIGHT
    row += 2

    # 毛利
    ws.cell(row=row, column=1, value="三、營業毛利").font = bold
    ws.cell(row=row, column=1).fill = _LIGHT_BLUE
    c = ws.cell(row=row, column=3, value=pl_data["gross_profit"])
    c.font = bold
    c.number_format = _MONEY_FMT
    c.alignment = _RIGHT
    row += 2

    # 營業費用
    ws.cell(row=row, column=1, value="四、營業費用").font = bold
    ws.cell(row=row, column=1).fill = _LIGHT_YELLOW
    row += 1
    for d in pl_data["expense_detail"]:
        ws.cell(row=row, column=1, value=f"  {d['name']}").font = _NORMAL_FONT
        c = ws.cell(row=row, column=2, value=d["amount"])
        c.font = _NORMAL_FONT
        c.number_format = _MONEY_FMT
        c.alignment = _RIGHT
        row += 1
    ws.cell(row=row, column=1, value="  營業費用合計").font = bold
    c = ws.cell(row=row, column=3, value=pl_data["expense"])
    c.font = bold
    c.number_format = _MONEY_FMT
    c.alignment = _RIGHT
    row += 2

    # 淨利
    ws.cell(row=row, column=1, value="五、營業淨利（損）").font = Font(
        name="微軟正黑體", size=12, bold=True)
    c = ws.cell(row=row, column=3, value=pl_data["net_income"])
    c.font = Font(name="微軟正黑體", size=12, bold=True,
                  color="008000" if pl_data["net_income"] >= 0 else "FF0000")
    c.number_format = _MONEY_FMT
    c.alignment = _RIGHT


def _write_balance_sheet(ws, bs_data: dict):
    """Sheet 7: 資產負債表"""
    ym = bs_data["year_month"]
    ws.merge_cells("A1:C1")
    ws["A1"] = f"資產負債表 — {ym}"
    ws["A1"].font = _TITLE_FONT

    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 16

    row = 3
    bold = Font(name="微軟正黑體", size=11, bold=True)

    # 資產
    ws.cell(row=row, column=1, value="一、資產").font = bold
    ws.cell(row=row, column=1).fill = _LIGHT_GREEN
    row += 1
    for d in bs_data["asset_detail"]:
        ws.cell(row=row, column=1, value=f"  {d['name']}").font = _NORMAL_FONT
        c = ws.cell(row=row, column=2, value=d["amount"])
        c.font = _NORMAL_FONT
        c.number_format = _MONEY_FMT
        c.alignment = _RIGHT
        row += 1
    ws.cell(row=row, column=1, value="  資產合計").font = bold
    c = ws.cell(row=row, column=3, value=bs_data["assets"])
    c.font = bold
    c.number_format = _MONEY_FMT
    c.alignment = _RIGHT
    row += 2

    # 負債
    ws.cell(row=row, column=1, value="二、負債").font = bold
    ws.cell(row=row, column=1).fill = _LIGHT_RED
    row += 1
    for d in bs_data["liability_detail"]:
        ws.cell(row=row, column=1, value=f"  {d['name']}").font = _NORMAL_FONT
        c = ws.cell(row=row, column=2, value=d["amount"])
        c.font = _NORMAL_FONT
        c.number_format = _MONEY_FMT
        c.alignment = _RIGHT
        row += 1
    ws.cell(row=row, column=1, value="  負債合計").font = bold
    c = ws.cell(row=row, column=3, value=bs_data["liabilities"])
    c.font = bold
    c.number_format = _MONEY_FMT
    c.alignment = _RIGHT
    row += 2

    # 權益
    ws.cell(row=row, column=1, value="三、權益").font = bold
    ws.cell(row=row, column=1).fill = _LIGHT_BLUE
    row += 1
    for d in bs_data["equity_detail"]:
        ws.cell(row=row, column=1, value=f"  {d['name']}").font = _NORMAL_FONT
        c = ws.cell(row=row, column=2, value=d["amount"])
        c.font = _NORMAL_FONT
        c.number_format = _MONEY_FMT
        c.alignment = _RIGHT
        row += 1
    ws.cell(row=row, column=1, value="  權益合計").font = bold
    c = ws.cell(row=row, column=3, value=bs_data["equity"])
    c.font = bold
    c.number_format = _MONEY_FMT
    c.alignment = _RIGHT
    row += 2

    # 平衡驗證
    status = "✅ 資產 = 負債 + 權益" if bs_data["balanced"] else "❌ 資產負債表不平衡！"
    ws.cell(row=row, column=1, value=status).font = Font(
        name="微軟正黑體", size=12, bold=True,
        color="008000" if bs_data["balanced"] else "FF0000")


def _write_general_ledger_sheet(ws, ledger_entries: list, year_month: str):
    """Sheet 8: 總分類帳"""
    ws.merge_cells("A1:G1")
    ws["A1"] = f"總分類帳 — {year_month}"
    ws["A1"].font = _TITLE_FONT

    headers = ["科目代碼", "科目名稱", "日期", "摘要", "借方", "貸方", "餘額"]
    widths = [10, 16, 12, 24, 12, 12, 14]

    for col_idx, (h, w) in enumerate(zip(headers, widths), 1):
        cell = ws.cell(row=3, column=col_idx, value=h)
        cell.font = _HEADER_FONT
        cell.fill = _BLUE_FILL
        cell.alignment = _CENTER
        cell.border = _THIN_BORDER
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    row = 4
    current_account = None
    running = 0

    for e in ledger_entries:
        acct = e.get("account_code", "")
        if acct != current_account:
            current_account = acct
            running = 0
            if row > 4:
                row += 1  # 空行分隔科目

        running += (e.get("debit", 0) or 0) - (e.get("credit", 0) or 0)

        ws.cell(row=row, column=1, value=acct).font = _NORMAL_FONT
        ws.cell(row=row, column=2, value=e.get("account_name", "")).font = _NORMAL_FONT
        ws.cell(row=row, column=3, value=e.get("entry_date", "")).font = _NORMAL_FONT
        ws.cell(row=row, column=4, value=e.get("description", "")).font = _NORMAL_FONT

        d = ws.cell(row=row, column=5, value=e.get("debit", 0) or "")
        d.font = _NORMAL_FONT
        if e.get("debit"):
            d.number_format = _MONEY_FMT
        d.alignment = _RIGHT

        c = ws.cell(row=row, column=6, value=e.get("credit", 0) or "")
        c.font = _NORMAL_FONT
        if e.get("credit"):
            c.number_format = _MONEY_FMT
        c.alignment = _RIGHT

        bal = ws.cell(row=row, column=7, value=running)
        bal.font = _NORMAL_FONT
        bal.number_format = _MONEY_FMT
        bal.alignment = _RIGHT

        for col in range(1, 8):
            ws.cell(row=row, column=col).border = _THIN_BORDER
        row += 1


def _update_monthly_accounting(year_month: str, excel_path: str,
                               journal_entries: list, income_rows: list,
                               stagings: list):
    """更新月度會計總表"""
    total_income = sum(r.get("amount", 0) for r in income_rows)
    total_expense = sum(s.get("total_amount", 0) for s in stagings)
    total_tax = sum(s.get("tax_amount", 0) for s in stagings)
    net_profit = total_income - total_expense

    sm.upsert_monthly_accounting(
        year_month,
        total_income=total_income,
        total_expense=total_expense,
        total_tax=total_tax,
        net_profit=net_profit,
        journal_count=len(journal_entries),
        excel_path=excel_path,
    )


# =====================================================================
# 9. 歸檔後一鍵會計（主入口）
# =====================================================================

def process_after_archive(staging_id: int) -> str:
    """歸檔後自動完成所有會計流程"""
    staging = sm.get_staging(staging_id)
    if not staging:
        return ""

    year_month = staging.get("year_month") or datetime.now().strftime("%Y-%m")

    try:
        entries = generate_journal_entries(staging_id)
        balance = verify_balance(staging_id)

        excel_path = None
        try:
            excel_path = generate_accounting_excel(year_month)
        except Exception as e:
            logger.warning(f"Excel generation failed (non-fatal): {e}")

        parts = [f"\n📊 會計：{len(entries)} 筆分錄"]
        if balance["balanced"]:
            parts.append("✅ 借貸平衡")
        else:
            parts.append(f"⚠️ 差額 ${balance['difference']:,.0f}")
        if excel_path:
            parts.append(f"📁 帳冊已更新")

        return " | ".join(parts)

    except Exception as e:
        logger.error(f"Accounting process failed for staging #{staging_id}: {e}")
        return f"\n⚠️ 會計處理異常：{e}"


# =====================================================================
# 10. 月度報表快捷
# =====================================================================

def get_monthly_report_text(year_month: str) -> str:
    """取得月度會計文字摘要（LINE 回覆用）"""
    summary = sm.get_journal_summary(year_month)
    accounting = sm.get_monthly_accounting(year_month)
    stagings = sm.get_stagings_by_month(year_month)

    lines = [f"📊 {year_month} 會計摘要", ""]

    # 支出
    total_expense = sum(s.get("total_amount", 0) for s in stagings)
    total_tax = sum(s.get("tax_amount", 0) for s in stagings)
    lines.append(f"📦 進貨支出：${total_expense:,.0f}（{len(stagings)} 筆）")
    lines.append(f"💰 進項稅額：${total_tax:,.0f}")

    # 收入
    income_rows = sm.get_income_summary(year_month)
    total_income = sum(r.get("amount", 0) for r in income_rows)
    if total_income > 0:
        lines.append(f"💵 營業收入：${total_income:,.0f}")
        lines.append(f"📈 淨利：${total_income - total_expense:,.0f}")

    # 分錄
    lines.append(f"\n📝 分錄：{summary['count']} 筆")
    if summary["count"] > 0:
        status = "✅ 借貸平衡" if summary["balanced"] else "❌ 不平衡"
        lines.append(f"   借方合計：${summary['total_debit']:,.0f}")
        lines.append(f"   貸方合計：${summary['total_credit']:,.0f}")
        lines.append(f"   {status}")

    # 薪資
    payrolls = sm.get_payroll_for_journal(year_month)
    if payrolls:
        total_salary = sum(p.get("gross_salary", 0) or 0 for p in payrolls)
        lines.append(f"\n👥 薪資支出：${total_salary:,.0f}（{len(payrolls)} 人）")

    # Excel 路徑
    if accounting and accounting.get("excel_path"):
        lines.append(f"\n📁 帳冊：會計帳冊/{year_month}/")

    # 結帳狀態
    if accounting and accounting.get("is_closed"):
        lines.append(f"🔒 本月已結帳")

    return "\n".join(lines)


# =====================================================================
# 11. 教育訓練文件生成
# =====================================================================

def generate_training_document(output_dir: str = None) -> str:
    """生成會計系統教育訓練 Word 文件

    Returns: 文件路徑
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not output_dir:
        output_dir = os.path.join(ACCOUNTING_DIR, "教育訓練")
    _ensure_dir(output_dir)

    doc = Document()

    # 設定預設字型
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微軟正黑體"
    font.size = Pt(11)

    # === 封面 ===
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("小膳會計系統", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("操作手冊與教育訓練教材", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"版本：v2.7.0\n")
    info.add_run(f"更新日期：{datetime.now().strftime('%Y-%m-%d')}\n")
    info.add_run("適用對象：內部會計人員 / 管理人員\n")
    info.add_run("系統類型：AI 輔助會計助理")

    doc.add_page_break()

    # === 目錄 ===
    doc.add_heading("目錄", level=1)
    toc_items = [
        "一、系統概述",
        "二、基本操作流程",
        "  2.1 進貨記帳（拍照即記帳）",
        "  2.2 收入登記",
        "  2.3 薪資處理",
        "  2.4 固定資產管理",
        "三、會計報表",
        "  3.1 試算表",
        "  3.2 損益表",
        "  3.3 資產負債表",
        "  3.4 總分類帳",
        "四、期末結帳流程",
        "五、營業稅申報（401）",
        "六、LINE Bot 指令總覽",
        "七、常見問題 FAQ",
        "八、需要人工補充的資料清單",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number" if not item.startswith("  ") else "List Bullet")

    doc.add_page_break()

    # === 一、系統概述 ===
    doc.add_heading("一、系統概述", level=1)
    doc.add_paragraph(
        "小膳會計系統是一套專為團膳業設計的 AI 會計助理。"
        "它結合了 LINE Bot 介面和 AI 辨識能力，能夠自動完成從拍照到記帳的完整流程。"
    )
    doc.add_paragraph("系統核心功能：")
    features = [
        "拍照即記帳：手機拍攝收據/發票 → AI 自動辨識 → 生成複式分錄",
        "完整複式簿記：每筆交易自動產生借方和貸方分錄，確保借貸平衡",
        "自動化報表：試算表、損益表、資產負債表、總分類帳一鍵生成",
        "薪資自動入帳：薪資計算結果自動轉為會計分錄",
        "營業稅彙總：自動計算進銷項稅額，輔助 401 申報",
        "期末結帳：一鍵完成收入費用結轉",
        "Excel 帳冊：自動生成 8 頁完整會計帳冊（存 GDrive）",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("遵循標準", level=2)
    doc.add_paragraph(
        "本系統遵循台灣中小企業會計準則（EAS），採用二級會計科目編碼。"
        "科目架構如下："
    )

    # 科目表
    table = doc.add_table(rows=8, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["代碼區間", "類別", "說明"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    acct_data = [
        ("1000-1999", "資產", "現金、應收帳款、存貨、固定資產"),
        ("2000-2999", "負債", "應付帳款、應付薪資、銷項稅額"),
        ("3000-3999", "權益", "資本額、累積盈虧"),
        ("4000-4999", "收入", "營業收入、其他收入"),
        ("5000-5999", "成本", "進貨、直接人工、製造費用"),
        ("6000-6999", "費用", "薪資、租金、水電、折舊"),
        ("", "合計", "共 50+ 科目，涵蓋團膳業所有會計需求"),
    ]
    for i, (code, cat, desc) in enumerate(acct_data):
        row = table.rows[i + 1]
        row.cells[0].text = code
        row.cells[1].text = cat
        row.cells[2].text = desc

    doc.add_page_break()

    # === 二、基本操作流程 ===
    doc.add_heading("二、基本操作流程", level=1)

    doc.add_heading("2.1 進貨記帳（拍照即記帳）", level=2)
    doc.add_paragraph("這是最常用的功能。操作步驟：")
    steps = [
        "步驟 1：在 LINE 群組中拍攝收據或發票照片",
        "步驟 2：系統自動 OCR 辨識，顯示辨識結果（供應商、金額、品項）",
        "步驟 3：確認資料正確後，按「確認歸檔」",
        "步驟 4：系統自動完成以下工作：",
        "  - 生成複式分錄（借：進貨科目 / 貸：現金）",
        "  - 如有稅額且可扣抵 → 自動記錄進項稅額",
        "  - 更新月度 Excel 帳冊",
        "  - 圖片歸檔到 GDrive",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Bullet" if not s.startswith("  ") else "List Bullet 2")

    doc.add_paragraph("")
    doc.add_paragraph("分錄範例：向「好鮮水產行」進貨 $5,000（含稅），可扣抵：")
    je_table = doc.add_table(rows=4, cols=4)
    je_table.style = "Light Grid Accent 1"
    je_headers = ["摘要", "科目", "借方", "貸方"]
    for i, h in enumerate(je_headers):
        je_table.rows[0].cells[i].text = h
    je_data = [
        ("進貨-好鮮水產行", "進貨—水產類 (5110)", "$4,762", ""),
        ("進項稅額-好鮮水產行", "進項稅額 (1150)", "$238", ""),
        ("付款-好鮮水產行", "現金 (1100)", "", "$5,000"),
    ]
    for i, (desc, acct, dr, cr) in enumerate(je_data):
        row = je_table.rows[i + 1]
        row.cells[0].text = desc
        row.cells[1].text = acct
        row.cells[2].text = dr
        row.cells[3].text = cr

    doc.add_heading("2.2 收入登記", level=2)
    doc.add_paragraph(
        "收入登記目前可透過 LINE 指令或管理介面操作。"
        "系統會自動計算銷項稅額（5%）並產生對應分錄。"
    )
    doc.add_paragraph("分錄範例：收到團膳收入 $105,000（含稅）：")
    inc_table = doc.add_table(rows=4, cols=4)
    inc_table.style = "Light Grid Accent 1"
    for i, h in enumerate(je_headers):
        inc_table.rows[0].cells[i].text = h
    inc_data = [
        ("收款-團膳收入", "現金 (1100)", "$105,000", ""),
        ("收入-團膳收入", "營業收入 (4100)", "", "$100,000"),
        ("銷項稅額-團膳收入", "銷項稅額 (2150)", "", "$5,000"),
    ]
    for i, (desc, acct, dr, cr) in enumerate(inc_data):
        row = inc_table.rows[i + 1]
        row.cells[0].text = desc
        row.cells[1].text = acct
        row.cells[2].text = dr
        row.cells[3].text = cr

    doc.add_heading("2.3 薪資處理", level=2)
    doc.add_paragraph("薪資處理流程：")
    salary_steps = [
        "1. 下載薪資表 Excel 範本（透過 LINE 指令：「薪資表」）",
        "2. 填寫每位員工的底薪、加班、獎金等資料",
        "3. 上傳填寫完成的 Excel 到 LINE 群組",
        "4. 系統自動計算勞健保、勞退、所得稅扣繳",
        "5. 確認後，系統自動產生薪資分錄：",
        "   借：薪資費用 / 加班費 / 伙食費 / 勞健保費用 / 勞退費用",
        "   貸：現金 / 代扣所得稅 / 應付勞健保 / 應付勞退",
    ]
    for s in salary_steps:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("2.4 固定資產管理", level=2)
    doc.add_paragraph(
        "團膳業常見固定資產包括冷藏設備、爐具、送餐車等。"
        "系統支援直線法折舊，每月自動計算折舊金額並產生分錄。"
    )

    doc.add_page_break()

    # === 三、會計報表 ===
    doc.add_heading("三、會計報表", level=1)
    doc.add_paragraph(
        "系統每次歸檔或手動觸發時，會自動生成完整的 Excel 帳冊，包含以下 8 個工作表：")

    sheets = [
        ("進貨日記帳", "記錄所有進貨明細（日期/供應商/品項/金額）"),
        ("月度費用彙總", "按供應商和按分類兩種角度彙總支出"),
        ("試算表", "各科目借貸合計，驗證整月借貸是否平衡"),
        ("分錄明細", "所有會計分錄的完整清單（含來源追蹤）"),
        ("收入明細", "所有營業收入記錄"),
        ("損益表", "營業收入 - 營業成本 - 營業費用 = 營業淨利"),
        ("資產負債表", "資產 = 負債 + 權益 的平衡驗證"),
        ("總分類帳", "按科目分類的所有交易明細（含逐筆餘額）"),
    ]

    sheet_table = doc.add_table(rows=len(sheets) + 1, cols=2)
    sheet_table.style = "Light Grid Accent 1"
    sheet_table.rows[0].cells[0].text = "工作表"
    sheet_table.rows[0].cells[1].text = "內容"
    for i, (name, desc) in enumerate(sheets):
        sheet_table.rows[i + 1].cells[0].text = name
        sheet_table.rows[i + 1].cells[1].text = desc

    doc.add_page_break()

    # === 四、期末結帳 ===
    doc.add_heading("四、期末結帳流程", level=1)
    doc.add_paragraph("每月底或隔月初進行期末結帳，步驟如下：")
    closing_steps = [
        "1. 確認所有進貨收據已拍照歸檔",
        "2. 確認本月收入已全部登記",
        "3. 確認薪資已計算並入帳",
        "4. 在 LINE 中輸入「結帳」或「結帳 2026-03」",
        "5. 系統自動執行：",
        "   a. 計算固定資產折舊",
        "   b. 結轉收入科目至「本期損益」",
        "   c. 結轉成本費用至「本期損益」",
        "   d. 標記該月為已結帳",
        "6. 系統回報結帳結果（收入/成本/費用/淨利）",
    ]
    for s in closing_steps:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph(
        "注意：結帳後該月的分錄原則上不應修改。"
        "如需修改，請先取消結帳再進行調整。"
    )

    # === 五、營業稅 ===
    doc.add_heading("五、營業稅申報（401）", level=1)
    doc.add_paragraph(
        "營業稅為雙月申報制（1-2月、3-4月...），"
        "每期在次月 15 日前申報。"
    )
    doc.add_paragraph("系統提供的輔助功能：")
    vat_features = [
        "自動彙總雙月份的進項稅額（來自進貨發票）",
        "自動彙總雙月份的銷項稅額（來自收入記錄）",
        "計算應納稅額或留抵稅額",
        "LINE 指令：「營業稅」或「401」即可查看摘要",
    ]
    for f in vat_features:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph(
        "重要提醒：目前系統產出的營業稅摘要為參考用，"
        "正式申報仍需由會計人員將資料填入國稅局 401 表單。"
        "未來版本將支援 401 媒體檔匯出。"
    )

    doc.add_page_break()

    # === 六、LINE Bot 指令 ===
    doc.add_heading("六、LINE Bot 指令總覽", level=1)

    commands = [
        ("拍照", "拍攝收據/發票照片即可記帳"),
        ("會計", "查看本月會計摘要"),
        ("會計 2026-03", "查看指定月份會計摘要"),
        ("更新帳冊", "手動重新生成 Excel 帳冊"),
        ("結帳", "執行本月期末結帳"),
        ("結帳 2026-03", "執行指定月份期末結帳"),
        ("營業稅", "查看當期營業稅摘要"),
        ("損益表", "查看本月損益表文字版"),
        ("薪資表", "下載薪資表 Excel 範本"),
        ("待處理", "查看待確認的進貨記錄"),
        ("報表", "開啟報表選單"),
        ("說明", "查看系統操作說明"),
    ]

    cmd_table = doc.add_table(rows=len(commands) + 1, cols=2)
    cmd_table.style = "Light Grid Accent 1"
    cmd_table.rows[0].cells[0].text = "指令"
    cmd_table.rows[0].cells[1].text = "功能說明"
    for i, (cmd, desc) in enumerate(commands):
        cmd_table.rows[i + 1].cells[0].text = cmd
        cmd_table.rows[i + 1].cells[1].text = desc

    doc.add_page_break()

    # === 七、FAQ ===
    doc.add_heading("七、常見問題 FAQ", level=1)

    faqs = [
        ("Q: 拍照辨識不正確怎麼辦？",
         "A: 系統顯示辨識結果後，你可以在確認前手動修改。輸入「修改」即可進入編輯模式。"),
        ("Q: 借貸不平衡怎麼辦？",
         "A: 系統設計上會自動確保借貸平衡（貸方 = 借方合計）。"
         "如果出現不平衡，請聯繫系統管理員檢查。"),
        ("Q: 已歸檔的資料可以刪除嗎？",
         "A: 可以，但會同時刪除對應的分錄。建議先匯出帳冊備份。"),
        ("Q: 系統會自動備份嗎？",
         "A: Excel 帳冊自動存到 GDrive（會計帳冊/YYYY-MM/）。"
         "資料庫需由管理員定期備份。"),
        ("Q: 可以一次補記多筆舊的收據嗎？",
         "A: 可以。拍照後系統會辨識日期，如日期有誤可手動修改。"),
        ("Q: 不同月份的帳可以同時處理嗎？",
         "A: 可以。系統會根據收據日期自動歸入正確月份。"),
    ]

    for q, a in faqs:
        doc.add_paragraph(q, style="List Bullet")
        doc.add_paragraph(a)
        doc.add_paragraph("")

    # === 八、需要的資料 ===
    doc.add_heading("八、需要人工補充的資料清單", level=1)
    doc.add_paragraph(
        "為了讓會計系統完整運作，以下資料需要人工定期補充：")

    needed = [
        ("每日必做", [
            "拍攝所有進貨收據/發票（這是最重要的！）",
        ]),
        ("每月必做", [
            "登記本月營業收入（團膳收入、便當收入等）",
            "填寫薪資表並上傳",
            "執行期末結帳",
        ]),
        ("每雙月必做", [
            "核對營業稅摘要，向國稅局申報 401",
        ]),
        ("年度必做", [
            "年底盤點存貨（食材庫存估值）",
            "確認固定資產清單是否有新增/報廢",
            "提供年度營所稅申報所需資料給會計師",
        ]),
        ("建置期需要（一次性）", [
            "員工資料建檔（姓名/身分證/底薪/勞保級距）",
            "固定資產清單（設備名稱/購入日期/金額/使用年限）",
            "期初餘額（現金/應收帳款/存貨等各科目的初始餘額）",
            "資本額",
        ]),
    ]

    for category, items in needed:
        doc.add_heading(category, level=2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # === 附錄 ===
    doc.add_heading("附錄：完整會計科目表", level=1)
    doc.add_paragraph("以下為系統使用的完整會計科目（依中小企業會計準則）：")

    coa = sm.get_chart_of_accounts()
    if coa:
        coa_table = doc.add_table(rows=len(coa) + 1, cols=4)
        coa_table.style = "Light Grid Accent 1"
        coa_headers = ["代碼", "名稱", "類別", "正常方向"]
        for i, h in enumerate(coa_headers):
            coa_table.rows[0].cells[i].text = h
        cat_names = {
            "asset": "資產", "liability": "負債", "equity": "權益",
            "revenue": "收入", "cost": "成本", "expense": "費用",
        }
        for i, a in enumerate(coa):
            row = coa_table.rows[i + 1]
            row.cells[0].text = a["code"]
            row.cells[1].text = a["name"]
            row.cells[2].text = cat_names.get(a["category"], a["category"])
            row.cells[3].text = "借方" if a["normal_side"] == "debit" else "貸方"

    filepath = os.path.join(output_dir, f"小膳會計系統_教育訓練手冊_{datetime.now().strftime('%Y%m%d')}.docx")
    doc.save(filepath)
    logger.info(f"Training document generated: {filepath}")
    return filepath
