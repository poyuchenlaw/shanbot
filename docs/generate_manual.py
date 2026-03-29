#!/usr/bin/env python3
"""
生成：小膳 AI 內帳系統 — 員工操作手冊 v3.0
輸出：小膳BOT_員工操作手冊_v3.0.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── 顏色常數 ──
DEEP_BLUE = RGBColor(0x1A, 0x3E, 0x6E)
ORANGE = RGBColor(0xE8, 0x6C, 0x00)
LIGHT_ORANGE_BG = "FFF3E0"
LIGHT_BLUE_BG = "EBF0F7"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)

# ── 字型（微軟正黑體優先，fallback Noto Sans CJK TC）──
FONT_NAME = "Noto Sans CJK TC"  # WSL 上有安裝

doc = Document()

# ════════════════════════════════════════════
# 全域樣式設定
# ════════════════════════════════════════════

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(11)
style.font.color.rgb = DARK_GRAY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
# 設定東亞字型
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

# 設定各級標題樣式
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = FONT_NAME
    heading_style.font.color.rgb = DEEP_BLUE
    heading_style.font.bold = True
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if level == 1:
        heading_style.font.size = Pt(22)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    else:
        heading_style.font.size = Pt(13)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, color_hex):
    """設定表格儲存格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_run(paragraph, text, bold=False, color=None, size=None, font_name=None):
    """添加格式化的 run"""
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    fn = font_name or FONT_NAME
    run.font.name = fn
    run.element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    return run


def add_styled_paragraph(text, style_name='Normal', bold=False, color=None,
                          size=None, alignment=None, space_before=None,
                          space_after=None):
    """添加帶樣式的段落"""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p


def add_orange_box(text, is_bold_title=None):
    """添加橘色提示框（用表格模擬）"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_ORANGE_BG)
    # 設定邊框為橘色
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="8" w:color="E86C00"/>'
        '  <w:left w:val="single" w:sz="18" w:color="E86C00"/>'
        '  <w:bottom w:val="single" w:sz="8" w:color="E86C00"/>'
        '  <w:right w:val="single" w:sz="8" w:color="E86C00"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    # 設定寬度
    table.columns[0].width = Cm(16)

    p = cell.paragraphs[0]
    if is_bold_title:
        add_run(p, is_bold_title, bold=True, color=ORANGE, size=Pt(12))
        add_run(p, "\n" + text, color=DARK_GRAY, size=Pt(11))
    else:
        add_run(p, text, color=DARK_GRAY, size=Pt(11))
    doc.add_paragraph()  # 間距


def add_blue_box(text, title=None):
    """添加淺藍提示框"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE_BG)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="8" w:color="1A3E6E"/>'
        '  <w:left w:val="single" w:sz="18" w:color="1A3E6E"/>'
        '  <w:bottom w:val="single" w:sz="8" w:color="1A3E6E"/>'
        '  <w:right w:val="single" w:sz="8" w:color="1A3E6E"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    table.columns[0].width = Cm(16)

    p = cell.paragraphs[0]
    if title:
        add_run(p, title, bold=True, color=DEEP_BLUE, size=Pt(12))
        add_run(p, "\n" + text, color=DARK_GRAY, size=Pt(11))
    else:
        add_run(p, text, color=DARK_GRAY, size=Pt(11))
    doc.add_paragraph()


def add_step(number, text, detail=None):
    """添加步驟說明"""
    p = doc.add_paragraph()
    add_run(p, f"Step {number}  ", bold=True, color=DEEP_BLUE, size=Pt(13))
    add_run(p, text, bold=True, size=Pt(12))
    if detail:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        add_run(p2, detail, color=MID_GRAY, size=Pt(10.5))
    return p


def add_qa(question, answer):
    """添加 Q&A"""
    p = doc.add_paragraph()
    add_run(p, "Q: ", bold=True, color=DEEP_BLUE, size=Pt(11.5))
    add_run(p, question, bold=True, size=Pt(11.5))

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.8)
    add_run(p2, "A: ", bold=True, color=ORANGE, size=Pt(11))
    add_run(p2, answer, size=Pt(11))
    # 分隔線
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(8)
    return p


def add_bullet(text, bold_prefix=None, indent_cm=1.0):
    """添加項目符號"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent_cm)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, color=DEEP_BLUE)
        add_run(p, text)
    else:
        add_run(p, text)
    return p


# ════════════════════════════════════════════
# 頁面設定
# ════════════════════════════════════════════

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── 頁首 ──
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = hp.add_run("小膳 AI 內帳系統 — 員工操作手冊")
run.font.name = FONT_NAME
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
run.font.size = Pt(9)
run.font.color.rgb = MID_GRAY
# 頁首底線
pPr = hp._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    '  <w:bottom w:val="single" w:sz="4" w:color="CCCCCC" w:space="4"/>'
    '</w:pBdr>'
)
pPr.append(pBdr)

# ── 頁尾（頁碼）──
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
# 插入頁碼域代碼
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
run2 = fp.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run2._r.append(instrText)
run3 = fp.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run3._r.append(fldChar2)
fp.runs[0].font.size = Pt(9)
fp.runs[0].font.color.rgb = MID_GRAY

# ════════════════════════════════════════════
# 封面頁
# ════════════════════════════════════════════

# 空白間距
for _ in range(6):
    doc.add_paragraph()

# 主標題
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "小膳 AI 內帳系統", bold=True, color=DEEP_BLUE, size=Pt(36))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "員工操作手冊", bold=True, color=DEEP_BLUE, size=Pt(30))

# 版本
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
add_run(p, "v3.0", bold=True, color=ORANGE, size=Pt(24))

# 副標題
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
add_run(p, "五公司多帳號版", color=MID_GRAY, size=Pt(16))

# 分隔線（用表格模擬）
doc.add_paragraph()
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, "1A3E6E")
cell.width = Cm(10)
table.columns[0].width = Cm(10)
p = cell.paragraphs[0]
p.text = ""  # 薄條
# 設定行高很矮
tr = table.rows[0]._tr
trPr = tr.get_or_add_trPr()
trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="40" w:hRule="exact"/>')
trPr.append(trHeight)

doc.add_paragraph()

# 五家公司列表
companies = [
    ("福利社", "升鼎商行"),
    ("王凱", "王凱食品有限公司"),
    ("台達2廠", "升鼎商行"),
    ("富燚", "富燚商行"),
    ("台達1廠", "升鼎商行"),
]
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "適用公司：", bold=True, color=DEEP_BLUE, size=Pt(12))
for name, company in companies:
    add_run(p, f"\n{name}（{company}）", color=MID_GRAY, size=Pt(11))

# 日期
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "2026 年 3 月", color=MID_GRAY, size=Pt(12))

# ── 分頁 ──
doc.add_page_break()

# ════════════════════════════════════════════
# 目錄頁
# ════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "目    錄", bold=True, color=DEEP_BLUE, size=Pt(24))
doc.add_paragraph()

toc_items = [
    ("一、", "系統簡介", ""),
    ("二、", "加入方式", ""),
    ("三、", "拍照記帳（最常用功能）", ""),
    ("", "    拍照步驟 Step 1-5", ""),
    ("", "    拍照技巧", ""),
    ("四、", "查看統計", ""),
    ("五、", "六宮格功能總覽", ""),
    ("六、", "常見問題", ""),
    ("七、", "注意事項", ""),
    ("", "附錄：五公司 LINE 帳號總表", ""),
]

for prefix, title, _ in toc_items:
    p = doc.add_paragraph()
    if prefix:
        add_run(p, prefix, bold=True, color=DEEP_BLUE, size=Pt(13))
        add_run(p, title, bold=True, size=Pt(13))
    else:
        add_run(p, title, color=MID_GRAY, size=Pt(11))
    p.paragraph_format.space_after = Pt(8)
    # 添加底部虛線
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="dotted" w:sz="4" w:color="CCCCCC" w:space="2"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

doc.add_page_break()

# ════════════════════════════════════════════
# 一、系統簡介
# ════════════════════════════════════════════

doc.add_heading('一、系統簡介', level=1)

add_blue_box(
    "小膳是一套用 LINE 拍照就能記帳的 AI 系統。\n"
    "員工只要用手機拍下收據、發票，AI 就會自動辨識內容、分類、歸檔。",
    title="一句話了解小膳"
)

add_styled_paragraph("核心特點：", bold=True, color=DEEP_BLUE, size=Pt(13))
add_bullet("每家公司都有自己專屬的 LINE 帳號", bold_prefix="帳號獨立 — ")
add_bullet("各公司的帳本完全分開，資料不會互相干擾", bold_prefix="資料隔離 — ")
add_bullet("員工只需要加入自己公司的 LINE 帳號就好", bold_prefix="簡單加入 — ")
add_bullet("所有資料自動備份到 Google 雲端硬碟", bold_prefix="自動備份 — ")

doc.add_paragraph()

# 五家公司帳號表格
add_styled_paragraph("五家公司 LINE 帳號一覽：", bold=True, color=DEEP_BLUE, size=Pt(13))

table = doc.add_table(rows=6, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# 表頭
headers = ["公司簡稱", "商號全名", "LINE 帳號名稱"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "1A3E6E")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, color=WHITE, size=Pt(11))

data = [
    ("福利社", "升鼎商行", "福利社AI內帳彙整事務網"),
    ("王凱", "王凱食品有限公司", "王凱食品AI內帳彙整事務網"),
    ("台達2廠", "升鼎商行", "台達2廠AI內帳彙整事務網"),
    ("富燚", "富燚商行", "富燚食品AI內帳彙整事務網"),
    ("台達1廠", "升鼎商行", "台達1廠AI內帳彙整事務網"),
]

for row_idx, (short, full, line_name) in enumerate(data, 1):
    for col_idx, val in enumerate([short, full, line_name]):
        cell = table.rows[row_idx].cells[col_idx]
        if row_idx % 2 == 0:
            set_cell_shading(cell, "F5F5F5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, val, size=Pt(10.5))

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════
# 二、加入方式
# ════════════════════════════════════════════

doc.add_heading('二、加入方式', level=1)

add_styled_paragraph("只需要三步，就能開始使用小膳：", size=Pt(12))

doc.add_paragraph()

# Step 1
add_step(1, "打開 LINE，搜尋自己公司的帳號",
         "點選 LINE 上方搜尋框，輸入「XX AI內帳彙整事務網」，XX 換成自己的公司名。\n"
         "例如：「福利社AI內帳彙整事務網」、「王凱食品AI內帳彙整事務網」")

# Step 2
add_step(2, "加入好友",
         "也可以掃描 QR Code 加入（QR Code 請向管理員索取）")

# QR Code 預留位置
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, "F5F5F5")
table.columns[0].width = Cm(6)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "\n\n[ QR Code 預留位置 ]\n\n", color=MID_GRAY, size=Pt(14))
doc.add_paragraph()

# Step 3
add_step(3, "看到功能選單（六宮格）就完成了！",
         "加入好友後，聊天室底部會出現六個功能按鈕，就代表設定成功。")

add_orange_box(
    "請確認加入的是自己公司的帳號，不要加錯！\n"
    "例如：福利社的員工要加「福利社AI內帳彙整事務網」，不要加到王凱的。",
    is_bold_title="⚠ 重要提醒"
)

doc.add_page_break()

# ════════════════════════════════════════════
# 三、拍照記帳（最常用功能）
# ════════════════════════════════════════════

doc.add_heading('三、拍照記帳（最常用功能）', level=1)

add_blue_box(
    "拍照記帳是最常用的功能。收到收據、發票，拿手機拍一張傳上去就好。",
    title="這是你最常用的功能！"
)

doc.add_paragraph()
doc.add_heading('操作步驟', level=2)

# Step 1
add_step(1, "拍照上傳",
         "點選聊天室底部的「📸 拍照記帳」按鈕，或直接在聊天室拍照 / 上傳照片。")

# Step 2
add_step(2, "等待 AI 辨識",
         "系統會自動辨識收據內容，大約需要 3～5 秒。\n"
         "辨識完成後，系統會回傳辨識結果讓你確認。")

# Step 3
p = doc.add_paragraph()
add_run(p, "Step 3  ", bold=True, color=DEEP_BLUE, size=Pt(13))
add_run(p, "確認辨識結果", bold=True, size=Pt(12))

# 確認結果表格
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# 表頭
for i, h in enumerate(["辨識結果", "你要回覆"]):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "1A3E6E")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, color=WHITE, size=Pt(11))

confirm_data = [
    ("✅ 正確", "回覆「ok」「好」「確認」皆可"),
    ("❌ 有錯誤", "回覆「修改」，跟著系統引導修正"),
    ("🔄 要重拍", "回覆「放棄」，此張不入帳"),
]

for row_idx, (status, action) in enumerate(confirm_data, 1):
    cell0 = table.rows[row_idx].cells[0]
    cell1 = table.rows[row_idx].cells[1]
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, status, bold=True, size=Pt(11))
    p1 = cell1.paragraphs[0]
    add_run(p1, action, size=Pt(11))

doc.add_paragraph()

# Step 4
add_step(4, "最終確認 & 補填資訊",
         "如果是菜市場收據（沒有發票的），系統會額外詢問經手人姓名。\n"
         "填寫後系統會自動產生經手人憑證。")

# Step 5
add_step(5, "自動歸檔",
         "確認完成後，系統會自動：\n"
         "  ➊ 重新命名檔案（日期_商家_金額）\n"
         "  ➋ 分類到正確的資料夾\n"
         "  ➌ 備份到 Google 雲端硬碟\n"
         "  你不需要做任何事，等系統回覆「已歸檔」就完成了！")

doc.add_paragraph()
doc.add_heading('拍照技巧', level=2)

add_orange_box(
    "① 光線充足，不要逆光\n"
    "② 收據平放桌面，不要用手拿著拍\n"
    "③ 四個角都要入鏡，不要切到邊\n"
    "④ 對焦清楚再按快門\n"
    "⑤ 如果收據太長，分兩段拍也可以",
    is_bold_title="📸 五個拍照要訣"
)

doc.add_page_break()

# ════════════════════════════════════════════
# 四、查看統計
# ════════════════════════════════════════════

doc.add_heading('四、查看統計', level=1)

add_styled_paragraph("想知道目前帳務狀況？只要在聊天室輸入關鍵字就好：", size=Pt(12))
doc.add_paragraph()

table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

for i, h in enumerate(["輸入", "功能", "說明"]):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "1A3E6E")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, color=WHITE, size=Pt(11))

stat_data = [
    ("「統計」", "查看本月收支", "顯示本月累計支出、分類明細"),
    ("「待處理」", "查看未確認單據", "列出還沒確認的拍照記帳，方便追蹤"),
]
for row_idx, (cmd, func, desc) in enumerate(stat_data, 1):
    for col_idx, val in enumerate([cmd, func, desc]):
        cell = table.rows[row_idx].cells[col_idx]
        p = cell.paragraphs[0]
        if col_idx == 0:
            add_run(p, val, bold=True, color=ORANGE, size=Pt(11))
        else:
            add_run(p, val, size=Pt(11))

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════
# 五、六宮格功能總覽
# ════════════════════════════════════════════

doc.add_heading('五、六宮格功能總覽', level=1)

add_styled_paragraph("加入好友後，聊天室底部會出現這六個功能按鈕：", size=Pt(12))
doc.add_paragraph()

features = [
    ("📸 拍照記帳", "拍收據、發票，AI 自動辨識入帳", "最常用！"),
    ("📁 財務資料", "上傳 Excel / PDF 等財務文件", ""),
    ("🛒 採購管理", "查看待確認單據、供應商資料、市場行情", ""),
    ("🍽️ 菜單企劃", "菜單排程、成本試算、菜色照片上傳", ""),
    ("📊 報表生成", "會計帳冊、四大報表、稅務匯出", ""),
    ("❓ 使用說明", "操作教學、常見問題查詢", ""),
]

table = doc.add_table(rows=7, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

for i, h in enumerate(["功能按鈕", "說明", "備註"]):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "1A3E6E")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, color=WHITE, size=Pt(11))

for row_idx, (name, desc, note) in enumerate(features, 1):
    cell0 = table.rows[row_idx].cells[0]
    cell1 = table.rows[row_idx].cells[1]
    cell2 = table.rows[row_idx].cells[2]

    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, name, bold=True, size=Pt(11))

    p1 = cell1.paragraphs[0]
    add_run(p1, desc, size=Pt(10.5))

    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if note:
        add_run(p2, note, bold=True, color=ORANGE, size=Pt(11))

    if row_idx == 1:
        set_cell_shading(cell0, LIGHT_ORANGE_BG)
        set_cell_shading(cell1, LIGHT_ORANGE_BG)
        set_cell_shading(cell2, LIGHT_ORANGE_BG)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════
# 六、常見問題
# ════════════════════════════════════════════

doc.add_heading('六、常見問題', level=1)

add_qa("拍錯照片怎麼辦？",
       "回覆「放棄」就好，不會入帳。照片也不會留在系統裡。")

add_qa("AI 辨識結果不對？",
       "回覆「修改」，系統會一步一步問你要改什麼（例如：金額、日期、商家名稱），照著改就好。")

add_qa("菜市場沒有發票怎麼辦？",
       "照拍！系統會自動判斷這是菜市場收據，會額外問你經手人是誰，然後自動產生經手人憑證。")

add_qa("重複上傳同一張照片會怎樣？",
       "系統會偵測到重複，會問你要「跳過」還是「另存一筆」。選跳過就好。")

add_qa("要看本月報表？",
       "在聊天室輸入「統計」就能看到本月收支摘要。\n"
       "要更詳細的報表，點選「📊 報表生成」按鈕。")

add_qa("要匯出稅務資料？",
       "點選「📊 報表生成」→ 選擇「匯出功能」→ 選擇「稅務申報檔」，系統會自動產生並下載。")

add_qa("系統一直沒回應？",
       "等 10 秒再試一次。如果還是沒回應，請聯繫管理員（小魚）。")

doc.add_page_break()

# ════════════════════════════════════════════
# 七、注意事項
# ════════════════════════════════════════════

doc.add_heading('七、注意事項', level=1)

notes = [
    ("不要傳錯帳號", "每家公司的 LINE 帳號是獨立的。福利社的收據要傳到福利社的帳號，不要傳到其他公司去。"),
    ("資料自動備份", "所有照片和帳務資料都會自動備份到 Google 雲端硬碟，不用擔心資料遺失。"),
    ("有問題找管理員", "操作上有任何問題，請找小魚或聯繫管理員處理。"),
    ("手機更新 LINE", "請保持 LINE 在最新版本，以免功能無法正常使用。"),
    ("保護帳號安全", "不要把 LINE 帳號的登入資訊分享給其他人。"),
]

for idx, (title, desc) in enumerate(notes, 1):
    p = doc.add_paragraph()
    add_run(p, f"  {idx}.  ", bold=True, color=DEEP_BLUE, size=Pt(13))
    add_run(p, title, bold=True, color=DEEP_BLUE, size=Pt(13))
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.2)
    add_run(p2, desc, size=Pt(11))
    p2.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

add_orange_box(
    "最重要的一件事：收據拿到就拍，不要囤！\n"
    "每天拍完當天的收據，帳務就不會亂。",
    is_bold_title="⭐ 養成好習慣"
)

doc.add_page_break()

# ════════════════════════════════════════════
# 附錄：五公司 LINE 帳號總表
# ════════════════════════════════════════════

doc.add_heading('附錄：五公司 LINE 帳號總表', level=1)

add_styled_paragraph("請找到自己公司，加入對應的 LINE 帳號：", size=Pt(12))
doc.add_paragraph()

table = doc.add_table(rows=6, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

headers = ["#", "公司簡稱", "商號全名", "LINE 帳號"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "1A3E6E")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, color=WHITE, size=Pt(11))

appendix_data = [
    ("1", "福利社", "升鼎商行", "福利社AI內帳彙整事務網"),
    ("2", "王凱", "王凱食品有限公司", "王凱食品AI內帳彙整事務網"),
    ("3", "台達2廠", "升鼎商行", "台達2廠AI內帳彙整事務網"),
    ("4", "富燚", "富燚商行", "富燚食品AI內帳彙整事務網"),
    ("5", "台達1廠", "升鼎商行", "台達1廠AI內帳彙整事務網"),
]

for row_idx, row_data in enumerate(appendix_data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        if row_idx % 2 == 0:
            set_cell_shading(cell, "F5F5F5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_idx == 3:
            add_run(p, val, bold=True, color=DEEP_BLUE, size=Pt(10.5))
        else:
            add_run(p, val, size=Pt(10.5))

# 設定第一欄寬度較窄
table.columns[0].width = Cm(1.2)
table.columns[1].width = Cm(3)
table.columns[2].width = Cm(5)
table.columns[3].width = Cm(7)

doc.add_paragraph()
doc.add_paragraph()

# QR Code 預留
for company_short, _, line_name in [
    ("福利社", "升鼎商行", "福利社AI內帳彙整事務網"),
    ("王凱", "王凱食品有限公司", "王凱食品AI內帳彙整事務網"),
    ("台達2廠", "升鼎商行", "台達2廠AI內帳彙整事務網"),
    ("富燚", "富燚商行", "富燚食品AI內帳彙整事務網"),
    ("台達1廠", "升鼎商行", "台達1廠AI內帳彙整事務網"),
]:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # QR code placeholder
    cell0 = table.cell(0, 0)
    set_cell_shading(cell0, "F5F5F5")
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, "\n[ QR Code ]\n", color=MID_GRAY, size=Pt(11))
    table.columns[0].width = Cm(4)

    cell1 = table.cell(0, 1)
    p1 = cell1.paragraphs[0]
    add_run(p1, f"\n{company_short}\n", bold=True, color=DEEP_BLUE, size=Pt(14))
    add_run(p1, f"LINE 搜尋：{line_name}", color=MID_GRAY, size=Pt(10))
    table.columns[1].width = Cm(10)

    # 間距
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════
# 儲存
# ════════════════════════════════════════════

output_path = "/home/simon/shanbot/docs/小膳BOT_員工操作手冊_v3.0.docx"
doc.save(output_path)
print(f"✅ 手冊已產生：{output_path}")
print(f"   檔案大小：{os.path.getsize(output_path) / 1024:.1f} KB")
