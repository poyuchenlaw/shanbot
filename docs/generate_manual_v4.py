# -*- coding: utf-8 -*-
"""生成小膳 AI 內帳系統 — 員工操作手冊 v4.0"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import shutil

def set_cell_shading(cell, color):
    """設定儲存格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    """設定表格格線"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_table_with_header(doc, headers, rows, col_widths=None):
    """新增帶表頭底色的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # 表頭
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
        set_cell_shading(cell, "4472C4")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # 資料列
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    
    return table

def add_heading1(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.size = Pt(16)
    return h

def add_heading2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(14)
    return h

def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(12)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1 + level * 0.8)
    return p

def create_manual():
    doc = Document()
    
    # 設定預設字型
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(12)
    
    # 頁首頁尾
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        # 頁首
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "小膳 AI 內帳系統 — 員工操作手冊 v4.0"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        # 頁尾（頁碼）
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = footer_para.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = footer_para.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)

    # ========== 封面頁 ==========
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("小膳 AI 內帳系統")
    run.font.size = Pt(28)
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("員工操作手冊 v4.0")
    run.font.size = Pt(22)
    run.bold = True
    
    doc.add_paragraph()
    
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("五公司多帳號版")
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    companies = doc.add_paragraph()
    companies.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = companies.add_run("適用公司：福利社（升鼎商行）｜王凱（王凱食品有限公司）\n台達2廠（升鼎商行）｜富燚（富燚商行）｜台達1廠（升鼎商行）")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("2026 年 4 月")
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    ver_para = doc.add_paragraph()
    ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_para.add_run("版本：v4.0\n製作：小膳 AI 系統管理團隊")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # ========== 目錄 ==========
    doc.add_page_break()
    add_heading1(doc, "目錄")
    
    toc_items = [
        "一、系統簡介",
        "二、加入方式",
        "三、六宮格功能總覽",
        "四、拍照記帳（詳細操作）",
        "五、財務文件上傳",
        "六、報表生成與意涵",
        "七、員工修改機制",
        "八、月底自動結帳",
        "九、如何讀懂財務分析報告",
        "十、常見問題（FAQ）",
        "十一、注意事項",
        "附錄 A：五公司 LINE 帳號總表",
        "附錄 B：報表對照表",
        "附錄 C：會計科目簡表",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    # ========== 一、系統簡介 ==========
    doc.add_page_break()
    add_heading1(doc, "一、系統簡介")
    
    add_heading2(doc, "小膳是什麼？")
    add_para(doc, "小膳是一套專門為團膳公司設計的 AI 內帳系統。只要用手機打開 LINE，拍一張收據照片，系統就會自動幫你辨識金額、品項、供應商，然後記入帳本。從拍照到記帳，全部自動完成，不用手動輸入任何數字。")
    
    add_para(doc, "整個系統的流程是這樣的：")
    add_para(doc, "")
    
    flow_para = doc.add_paragraph()
    flow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = flow_para.add_run("拍照上傳 → AI辨識 → 確認入帳 → 自動分錄 → 月報表 → 財務分析")
    run.font.size = Pt(13)
    run.bold = True
    
    add_para(doc, "")
    
    add_heading2(doc, "五家公司獨立帳號")
    add_para(doc, "我們的五家公司各自有獨立的 LINE 帳號，資料完全隔離、互不干擾。拍照的時候要注意傳到正確的帳號，不要傳錯家喔！")
    
    add_table_with_header(doc, 
        ["公司簡稱", "公司全名", "統一編號"],
        [
            ["福利社", "升鼎商行", "81410187"],
            ["王凱", "王凱食品有限公司", "90438334"],
            ["台達2廠", "升鼎商行", "81410187"],
            ["富燚", "富燚商行", "00281384"],
            ["台達1廠", "升鼎商行", "81410187"],
        ]
    )
    
    add_para(doc, "")
    add_heading2(doc, "v4.0 新增功能")
    add_bullet(doc, "財務報表自動生成（資產負債表、損益表、現金流量表、權益變動表）")
    add_bullet(doc, "月底自動結帳（每月最後一天晚上 8 點自動執行）")
    add_bullet(doc, "財務分析報告（毛利率、成本率、趨勢分析、風險提醒）")
    add_bullet(doc, "自動稽核（6 項檢查，確保帳務正確）")
    add_bullet(doc, "財務文件上傳與智能分類")

    # ========== 二、加入方式 ==========
    doc.add_page_break()
    add_heading1(doc, "二、加入方式")
    
    add_para(doc, "加入步驟很簡單，三步就搞定：", bold=True)
    add_para(doc, "")
    
    add_para(doc, "Step 1：搜尋 LINE 帳號", bold=True)
    add_para(doc, "打開 LINE → 點上方搜尋 → 輸入你負責公司的帳號名稱（例如「福利社 AI內帳彙整事務網」）。每家公司的帳號名稱格式都是「XX AI內帳彙整事務網」。")
    add_para(doc, "")
    
    add_para(doc, "Step 2：加好友", bold=True)
    add_para(doc, "找到帳號後，點「加入好友」。如果公司有提供 QR Code，直接掃描更快。")
    add_para(doc, "")
    
    add_para(doc, "Step 3：進入系統", bold=True)
    add_para(doc, "加好友後，點進對話視窗，你會看到下方出現「六宮格」選單，就代表成功了！可以開始使用。")
    add_para(doc, "")
    
    add_para(doc, "💡 小提醒：如果你負責多家公司的記帳，要分別加入每家公司的帳號喔。每個帳號的資料是完全獨立的。")

    # ========== 三、六宮格功能總覽 ==========
    doc.add_page_break()
    add_heading1(doc, "三、六宮格功能總覽")
    
    add_para(doc, "打開 LINE 對話視窗，下方會看到六個功能格子。以下是每個格子的用途：")
    add_para(doc, "")
    
    add_table_with_header(doc,
        ["格子", "名稱", "功能說明"],
        [
            ["1", "📸 拍照記帳", "拍收據照片，AI 自動辨識入帳。這是最常用的功能！"],
            ["2", "📁 財務資料", "上傳 Excel/PDF 文件，系統自動分類歸檔"],
            ["3", "🛒 採購管理", "查看待處理單據、供應商資訊、市場行情"],
            ["4", "🍽️ 菜單企劃", "菜單排程、菜色圖片管理、食材成本試算"],
            ["5", "📊 報表生成", "產出四大財務報表 + 匯出功能"],
            ["6", "❓ 使用說明", "操作指南、常見問題 FAQ"],
        ]
    )
    
    add_para(doc, "")
    add_para(doc, "日常最常用的就是第 1 格「拍照記帳」，拍完收據就等著 AI 辨識結果，確認後就記好帳了。月底要看報表的話，點第 5 格就可以產出。")

    # ========== 四、拍照記帳（詳細操作） ==========
    doc.add_page_break()
    add_heading1(doc, "四、拍照記帳（詳細操作）")
    
    add_para(doc, "這是最核心、也是你每天都會用到的功能。以下是完整的操作流程：")
    add_para(doc, "")
    
    add_para(doc, "Step 1：開始拍照", bold=True)
    add_para(doc, "方法一：點下方六宮格的「📸 拍照記帳」按鈕")
    add_para(doc, "方法二：直接在對話視窗中傳送照片（更快！）")
    add_para(doc, "")
    
    add_para(doc, "Step 2：等待 AI 辨識", bold=True)
    add_para(doc, "照片送出後，大約 3～5 秒系統就會回覆辨識結果。結果會顯示：供應商名稱、日期、品項列表、金額合計、稅額等資訊。")
    add_para(doc, "")
    
    add_para(doc, "辨識信心度說明：", bold=True)
    add_table_with_header(doc,
        ["信心度", "符號", "意思", "你要做什麼"],
        [
            ["高", "🟢 綠色", "系統很有把握，辨識正確", "直接確認就好"],
            ["中", "🟡 黃色", "系統不太確定某些欄位", "請仔細看一下金額和品項對不對"],
            ["低", "🔴 紅色", "照片可能模糊或格式特殊", "建議重拍一張"],
        ]
    )
    add_para(doc, "")
    
    add_para(doc, "Step 3：確認辨識結果", bold=True)
    add_para(doc, "看完結果後，下方會出現三個按鈕：")
    add_bullet(doc, "✅ 確認：資料正確，直接入帳")
    add_bullet(doc, "✏️ 修改：有地方需要改（金額、日期等）")
    add_bullet(doc, "🗑️ 捨棄：拍錯了或重複了，不要這筆")
    add_para(doc, "")
    
    add_para(doc, "Step 4：菜市場收據額外步驟", bold=True)
    add_para(doc, "如果是菜市場（傳統市場）買的東西，系統會額外問你「經手人是誰？」，請輸入去買菜那個人的名字。這是為了內控和追蹤用的。")
    add_para(doc, "")
    
    add_para(doc, "Step 5：自動歸檔", bold=True)
    add_para(doc, "確認後系統會自動：")
    add_bullet(doc, "幫照片重新命名（格式：日期_供應商_金額）")
    add_bullet(doc, "上傳到 Google 雲端硬碟的對應資料夾")
    add_bullet(doc, "產生會計分錄記錄")
    add_para(doc, "")
    
    add_heading2(doc, "拍照技巧（很重要！）")
    add_para(doc, "照片品質直接影響 AI 辨識的準確度，請注意以下幾點：")
    add_bullet(doc, "💡 光線充足：不要在太暗的地方拍，自然光最好")
    add_bullet(doc, "📐 平放桌面：把收據攤平放在桌上，不要用手拿著拍")
    add_bullet(doc, "🔍 對焦清晰：確認畫面是清楚的，文字看得見")
    add_bullet(doc, "📷 四角入鏡：整張收據要完整出現在畫面中")
    add_bullet(doc, "✋ 穩定拍攝：拿穩手機，不要晃動")
    add_para(doc, "")
    
    add_heading2(doc, "背後發生了什麼？（了解即可）")
    add_para(doc, "你按下確認後，系統在背後會自動做這些事：")
    add_bullet(doc, "自動產生會計分錄（例如：借方「進貨」／貸方「現金」）")
    add_bullet(doc, "自動判斷這張發票可不可以扣抵營業稅（有統編的才可以）")
    add_bullet(doc, "自動歸類到正確的供應商資料夾")
    add_bullet(doc, "更新本月的採購統計")
    add_para(doc, "你不需要懂會計，系統全部幫你處理好。你只要確認金額和品項正確就行了！")

    # ========== 五、財務文件上傳 ==========
    doc.add_page_break()
    add_heading1(doc, "五、財務文件上傳")
    
    add_para(doc, "除了拍收據照片，你也可以上傳 Excel 或 PDF 文件到系統中，系統會自動讀取內容並分類歸檔。")
    add_para(doc, "")
    
    add_heading2(doc, "支援格式")
    add_bullet(doc, "Excel 檔案（.xlsx）")
    add_bullet(doc, "PDF 文件（.pdf）")
    add_para(doc, "")
    
    add_heading2(doc, "上傳流程")
    add_para(doc, "Step 1：點「📁 財務資料」或直接傳檔案到對話視窗")
    add_para(doc, "Step 2：系統自動讀取檔案內容（約 5～10 秒）")
    add_para(doc, "Step 3：系統判斷這份文件屬於哪個分類")
    add_para(doc, "Step 4：顯示分類結果，你可以「確認」或「修改分類」")
    add_para(doc, "Step 5：確認後自動歸檔到 Google 雲端硬碟對應資料夾")
    add_para(doc, "")
    
    add_heading2(doc, "八大循環分類表")
    add_para(doc, "系統會把文件自動歸到以下八大分類之一：")
    add_para(doc, "")
    
    add_table_with_header(doc,
        ["分類", "說明", "常見例子"],
        [
            ["收入循環", "團膳合約、收入明細相關", "月營收表、客戶合約"],
            ["支出循環", "採購、付款相關", "進貨單、付款明細"],
            ["人力資源", "薪資、勞健保相關", "薪資表、勞保名冊"],
            ["生產循環", "廚房作業、製程相關", "產量報表、生產紀錄"],
            ["融資循環", "借貸、銀行往來相關", "貸款合約、銀行對帳單"],
            ["投資循環", "設備購入、投資相關", "設備報價單、投資明細"],
            ["固定資產", "廚具、車輛、器材相關", "折舊表、資產清冊"],
            ["一般循環", "租約、其他合約", "場地租約、保險單"],
        ]
    )
    add_para(doc, "")
    add_para(doc, "💡 如果系統判斷錯分類，你可以在確認前修改。修改後系統會學習，下次遇到類似的文件就會分對。")

    # ========== 六、報表生成與意涵 ==========
    doc.add_page_break()
    add_heading1(doc, "六、報表生成與意涵")
    
    add_para(doc, "系統可以自動產出 9 種報表。以下用白話文說明每種報表是什麼意思、什麼時候要看：")
    add_para(doc, "")
    
    add_para(doc, "1. 月報表", bold=True)
    add_para(doc, "這是本月所有採購的彙總表。看一眼就知道這個月總共花了多少錢、買了幾筆。老闆最常看這張，了解整體花費狀況。")
    add_para(doc, "")
    
    add_para(doc, "2. 採購報告", bold=True)
    add_para(doc, "每筆採購的品項明細加上分類統計。可以看出哪一類食材花最多錢（例如肉類、蔬菜、調味料），幫助控制成本。")
    add_para(doc, "")
    
    add_para(doc, "3. 會計帳冊（8 頁）", bold=True)
    add_para(doc, "這是正式的會計記錄，總共有 8 頁，包括：進貨日記帳、費用彙總、試算表、分錄明細、收入記錄、損益表、資產負債表、總分類帳。主要是給會計師或記帳士看的，一般員工不需要細看。")
    add_para(doc, "")
    
    add_para(doc, "4. 資產負債表", bold=True)
    add_para(doc, "白話就是：公司現在有多少資產（現金、存貨、設備）、欠多少錢（應付帳款、貸款）、淨值多少。像是公司的「健康檢查報告」。")
    add_para(doc, "")
    
    add_para(doc, "5. 損益表", bold=True)
    add_para(doc, "白話就是：這個月賺多少錢、花了多少錢、最後有沒有賺。是老闆判斷公司經營狀況最重要的報表。")
    add_para(doc, "")
    
    add_para(doc, "6. 現金流量表", bold=True)
    add_para(doc, "錢從哪裡來（收到團膳款項）、花到哪裡去（買食材、付薪水）。可以看出公司現金是不是夠用。")
    add_para(doc, "")
    
    add_para(doc, "7. 權益變動表", bold=True)
    add_para(doc, "老闆的錢（業主權益）這個月怎麼變化。有投入新資金、或是賺了錢增加權益，都會記錄在這裡。")
    add_para(doc, "")
    
    add_para(doc, "8. 稽核報告", bold=True)
    add_para(doc, "系統自動檢查帳有沒有做錯。總共有 6 項檢查（借貸平衡、稅額覈對、日期連續性等），全部打勾 ✅ 就表示帳務正確。如果有紅色 ❌，代表有問題需要處理。")
    add_para(doc, "")
    
    add_para(doc, "9. 財務分析報告", bold=True)
    add_para(doc, "這是最高層級的報告，包含經營指標（毛利率、成本率）、趨勢分析、風險提醒、改進建議。幫助老闆做經營決策。")
    add_para(doc, "")
    
    add_para(doc, "💡 怎麼產出報表：點六宮格的「📊 報表生成」→ 選擇你要的報表類型 → 選月份 → 系統自動產出並提供下載連結。")

    # ========== 七、員工修改機制 ==========
    doc.add_page_break()
    add_heading1(doc, "七、員工修改機制")
    
    add_heading2(doc, "什麼可以改？")
    add_para(doc, "在按下「確認」之前，你可以修改以下欄位：")
    add_bullet(doc, "金額（發現 AI 辨識錯的時候）")
    add_bullet(doc, "供應商名稱")
    add_bullet(doc, "日期")
    add_bullet(doc, "食材分類")
    add_bullet(doc, "發票號碼")
    add_para(doc, "")
    
    add_heading2(doc, "怎麼改？")
    add_para(doc, "Step 1：在辨識結果出來後，按「✏️ 修改」按鈕")
    add_para(doc, "Step 2：系統會一個一個問你要改哪裡")
    add_para(doc, "Step 3：輸入正確的資料")
    add_para(doc, "Step 4：全部改完後，系統重新顯示結果讓你再次確認")
    add_para(doc, "")
    
    add_heading2(doc, "確認後還能改嗎？")
    add_para(doc, "已經按下「確認」的單據，員工自己不能修改。如果發現有錯，請聯繫管理員（小魚）處理。")
    add_para(doc, "")
    
    add_heading2(doc, "系統怎麼記錄修改？")
    add_para(doc, "每一次修改都會自動記錄：修改時間、修改者、修改前後的內容。管理員可以追溯所有修改歷史，所以請誠實填寫正確資料。")
    add_para(doc, "")
    
    add_para(doc, "⚠️ 重要提醒：月底結帳後的資料不能再修改。所以請在月底之前確認所有單據的資料都是正確的！", bold=True)

    # ========== 八、月底自動結帳 ==========
    doc.add_page_break()
    add_heading1(doc, "八、月底自動結帳")
    
    add_para(doc, "系統會在每月最後一天晚上 8 點自動執行月底結帳程序。以下是系統會自動做的事情：")
    add_para(doc, "")
    
    add_para(doc, "自動執行流程：", bold=True)
    add_bullet(doc, "① 確認所有待處理單據（如果有沒確認的，會自動提醒）")
    add_bullet(doc, "② 生成全部 9 種報表")
    add_bullet(doc, "③ 執行自動稽核（6 項檢查）")
    add_bullet(doc, "④ 期末結帳（結轉本月損益到權益）")
    add_bullet(doc, "⑤ 產出財務分析報告")
    add_bullet(doc, "⑥ 全部檔案歸檔到 Google 雲端硬碟")
    add_bullet(doc, "⑦ 推送月度摘要到 LINE 群組")
    add_para(doc, "")
    
    add_para(doc, "員工在月底前需要做的事：", bold=True)
    add_bullet(doc, "✅ 把所有收據照片傳完（不要留到下個月才傳）")
    add_bullet(doc, "✅ 確認所有待處理的單據（不要留 pending 狀態）")
    add_bullet(doc, "✅ 檢查有沒有漏拍的收據")
    add_bullet(doc, "✅ 如果發現有錯的，趕快修改或通知管理員")
    add_para(doc, "")
    add_para(doc, "⚠️ 月底結帳後，當月的資料就鎖定了，不能再修改。務必在每月最後一天晚上 8 點之前完成所有動作！", bold=True)

    # ========== 九、如何讀懂財務分析報告 ==========
    doc.add_page_break()
    add_heading1(doc, "九、如何讀懂財務分析報告")
    
    add_para(doc, "每個月系統會自動產出一份財務分析報告，裡面有些專業術語，以下用白話文幫你解釋：")
    add_para(doc, "")
    
    add_para(doc, "毛利率", bold=True)
    add_para(doc, "意思：營收（收到的團膳款）扣掉食材成本後，剩下多少比例。")
    add_para(doc, "白話：每收 100 元，扣掉買菜買肉的錢後還剩多少。")
    add_para(doc, "參考標準：30% 以上算健康，低於 20% 就要注意成本控制了。")
    add_para(doc, "")
    
    add_para(doc, "成本率", bold=True)
    add_para(doc, "意思：食材成本佔營收的比例。")
    add_para(doc, "白話：每收 100 元要花多少錢去買食材。越低代表成本控制越好。")
    add_para(doc, "參考標準：團膳業一般在 55%～70% 之間。")
    add_para(doc, "")
    
    add_para(doc, "供應商集中度", bold=True)
    add_para(doc, "意思：最大的供應商佔你採購金額的比例。")
    add_para(doc, "白話：是不是太依賴某一家供應商在買東西。")
    add_para(doc, "參考標準：單一供應商超過 40% 就要注意。如果那家出問題，我們會很麻煩。")
    add_para(doc, "")
    
    add_para(doc, "環比趨勢", bold=True)
    add_para(doc, "意思：跟上個月比較，數字是上升還是下降。")
    add_para(doc, "白話：這個月有沒有比上個月好或差。")
    add_para(doc, "")
    
    add_para(doc, "風險等級說明：", bold=True)
    add_table_with_header(doc,
        ["等級", "符號", "意思", "該怎麼做"],
        [
            ["高風險", "🔴 紅色", "有嚴重問題需要馬上處理", "立即通知管理員"],
            ["中風險", "🟡 黃色", "需要注意、持續觀察", "跟管理員討論是否要調整"],
            ["低風險", "🟢 綠色", "狀況正常，可以觀察", "維持目前做法就好"],
        ]
    )
    add_para(doc, "")
    add_para(doc, "💡 財務分析報告的建議是系統自動產生的，提供參考方向。實際要不要調整，還是由老闆和管理員決定。")

    # ========== 十、常見問題（FAQ） ==========
    doc.add_page_break()
    add_heading1(doc, "十、常見問題（FAQ）")
    
    faqs = [
        ("Q: 照片拍不清楚怎麼辦？", "A: 重拍一張就好。找光線充足的地方，把收據攤平在桌上，手機拿穩，確認對焦後再拍。"),
        ("Q: 系統辨識錯了怎麼辦？", "A: 按「✏️ 修改」按鈕，把錯的地方改正確。如果已經確認了才發現錯誤，聯繫管理員處理。"),
        ("Q: 一張收據有很多品項，AI 會全部辨識嗎？", "A: 會的。系統會把收據上所有品項和金額都列出來。如果有漏的，可以手動修改補上。"),
        ("Q: 可以一次傳多張照片嗎？", "A: 可以，一張一張傳就好。系統會依序處理，每張都會單獨回覆辨識結果。"),
        ("Q: 同一張收據不小心傳了兩次怎麼辦？", "A: 系統有重複偵測功能，會提醒你這張可能是重複的。如果真的重複了，選「捨棄」就好。"),
        ("Q: 月報表在哪裡看？", "A: 點「📊 報表生成」→ 匯出功能 → 月報表。也可以到 Google 雲端硬碟的「月報表」資料夾直接找。"),
        ("Q: 什麼是稽核報告？", "A: 系統自動檢查帳有沒有做錯的報告。如果 6 項檢查全部打勾 ✅，就代表帳務正確沒問題。"),
        ("Q: 財務分析報告的建議一定要照做嗎？", "A: 不一定。建議是系統自動產生的參考意見，實際決定還是由老闆和管理員判斷。"),
        ("Q: 上傳的文件跑到哪裡去了？", "A: 系統會自動分類到 Google 雲端硬碟對應的資料夾。點「📁 財務資料」→「查看本月文件」可以看到所有已上傳的文件。"),
        ("Q: 為什麼有些收據 AI 辨識不了？", "A: 可能是手寫的太潦草、照片太模糊、或是格式很特殊。遇到這種情況，系統會顯示 🔴 紅色提示，建議重拍或手動輸入。"),
        ("Q: 假日也可以用嗎？", "A: 可以。系統 24 小時運作，隨時都可以拍照上傳。"),
        ("Q: 手機沒網路的時候怎麼辦？", "A: 先把收據拍照存在手機裡，等有網路的時候再傳到 LINE。"),
    ]
    
    for q, a in faqs:
        add_para(doc, q, bold=True)
        add_para(doc, a)
        add_para(doc, "")

    # ========== 十一、注意事項 ==========
    doc.add_page_break()
    add_heading1(doc, "十一、注意事項")
    
    add_para(doc, "使用小膳系統，請特別注意以下幾點：")
    add_para(doc, "")
    
    items = [
        ("1. 不要傳錯帳號", "五家公司各有獨立的 LINE 帳號，資料完全隔離。傳照片前先確認你打開的是正確的帳號。"),
        ("2. 資料自動備份", "所有照片和文件都會自動備份到 Google 雲端硬碟，不用擔心資料遺失。"),
        ("3. 月底前把所有收據傳完", "每月最後一天晚上 8 點系統會自動結帳。在這之前要把當月所有收據都傳完並確認。"),
        ("4. 已確認的單據不要重複傳", "系統雖然有重複偵測功能，但為了避免混淆，請不要重複傳送已確認過的單據。"),
        ("5. 手機保持 LINE 最新版本", "LINE 版本太舊可能會影響功能正常使用。請定期更新 LINE App。"),
        ("6. 有問題找管理員（小魚）", "操作上有任何問題或發現帳務錯誤，請直接聯繫管理員小魚處理。"),
    ]
    
    for title, desc in items:
        add_para(doc, title, bold=True)
        add_para(doc, desc)
        add_para(doc, "")

    # ========== 附錄 A ==========
    doc.add_page_break()
    add_heading1(doc, "附錄 A：五公司 LINE 帳號總表")
    
    add_para(doc, "以下是五家公司的 LINE 帳號對照表：")
    add_para(doc, "")
    
    add_table_with_header(doc,
        ["公司簡稱", "LINE 帳號名稱", "公司全名", "統一編號"],
        [
            ["福利社", "福利社 AI內帳彙整事務網", "升鼎商行", "81410187"],
            ["王凱", "王凱 AI內帳彙整事務網", "王凱食品有限公司", "90438334"],
            ["台達2廠", "台達2廠 AI內帳彙整事務網", "升鼎商行", "81410187"],
            ["富燚", "富燚 AI內帳彙整事務網", "富燚商行", "00281384"],
            ["台達1廠", "台達1廠 AI內帳彙整事務網", "升鼎商行", "81410187"],
        ]
    )
    add_para(doc, "")
    add_para(doc, "💡 搜尋時如果找不到，試試只打公司簡稱（例如「福利社」或「王凱」）。")

    # ========== 附錄 B ==========
    doc.add_page_break()
    add_heading1(doc, "附錄 B：報表對照表")
    
    add_para(doc, "系統可產出的 9 種報表一覽：")
    add_para(doc, "")
    
    add_table_with_header(doc,
        ["報表名稱", "用途", "GDrive 位置", "產出時機"],
        [
            ["月報表", "月度採購彙總", "月報表/", "隨時產出 / 月底自動"],
            ["採購報告", "品項明細 + 分類統計", "採購單據/", "隨時產出 / 月底自動"],
            ["會計帳冊", "正式帳簿（8頁）", "會計資料/", "月底自動"],
            ["資產負債表", "公司財務狀況", "財務報表/", "隨時產出 / 月底自動"],
            ["損益表", "經營績效", "財務報表/", "隨時產出 / 月底自動"],
            ["現金流量表", "現金變化", "財務報表/", "隨時產出 / 月底自動"],
            ["權益變動表", "權益變化", "財務報表/", "隨時產出 / 月底自動"],
            ["稽核報告", "帳務正確性檢查", "會計資料/", "月底自動"],
            ["財務分析報告", "經營分析 + 改進建議", "會計資料/", "月底自動"],
        ]
    )
    add_para(doc, "")
    add_para(doc, "💡「隨時產出」代表你可以隨時到「📊 報表生成」去產出最新版本。「月底自動」代表每月底系統會自動生成並歸檔。")

    # ========== 附錄 C ==========
    doc.add_page_break()
    add_heading1(doc, "附錄 C：會計科目簡表")
    
    add_para(doc, "系統會自動產生會計分錄，以下是你可能會看到的科目名稱和它們的白話意思：")
    add_para(doc, "")
    
    add_table_with_header(doc,
        ["代碼", "科目名稱", "白話意思"],
        [
            ["1100", "現金", "我們付出去或收到的錢"],
            ["1150", "進項稅額", "買東西時付的稅（有統編的發票可以抵扣）"],
            ["4100", "營業收入", "團膳服務賺的錢（客戶付的餐費）"],
            ["5110", "進貨", "買食材花的錢（肉、菜、調味料等）"],
            ["6110", "薪資費用", "付給員工的薪水"],
            ["2150", "銷項稅額", "收入要繳給國稅局的稅"],
        ]
    )
    add_para(doc, "")
    add_para(doc, "💡 你不需要記住這些代碼，系統會全部自動處理。這裡只是讓你看到分錄時，知道那些數字代表什麼意思。")
    
    # ========== 文件尾 ==========
    doc.add_paragraph()
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("— 手冊結束 —")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    end_para2 = doc.add_paragraph()
    end_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para2.add_run("有任何問題請聯繫管理員（小魚）\n版本 v4.0 ｜ 2026 年 4 月")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    # 儲存
    path1 = "/home/simon/shanbot/docs/小膳BOT_員工操作手冊_v4.0.docx"
    path2 = "/mnt/h/小魚資料/團膳公司資料/小膳BOT_員工操作手冊_v4.0.docx"
    
    doc.save(path1)
    print(f"✅ 已儲存到: {path1}")
    
    shutil.copy2(path1, path2)
    print(f"✅ 已複製到: {path2}")
    
    print("完成！")

if __name__ == "__main__":
    create_manual()
