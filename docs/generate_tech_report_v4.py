"""生成小膳 Bot 技術開發報告 v4.0 Word 文件"""
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ======================================================================
# 樣式工具
# ======================================================================

def set_cell_shading(cell, color):
    """設定表格儲存格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_with_header(doc, headers, rows, col_widths=None):
    """建立有格線的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2E7D32")

    # Rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_code_block(doc, code_text):
    """加入等寬字型的程式碼區塊"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

def add_heading(doc, text, level=1):
    """加入標題"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14) if level == 1 else Pt(12)
    return h

# ======================================================================
# 主要生成邏輯
# ======================================================================

def generate_report():
    doc = Document()

    # -- 頁面設定 --
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # -- 頁首 --
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "小膳 Bot 技術開發報告 v4.0"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # -- 頁尾 --
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = "機密文件"
    for run in fp.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # 頁碼 (使用 XML field)
    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp2.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = fp2.add_run(' PAGE ')
    run2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    run3 = fp2.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)

    # ===================================================================
    # 封面頁
    # ===================================================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("小膳 Bot")
    run.font.size = Pt(36)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("技術開發報告 v4.0")
    run.font.size = Pt(24)
    run.bold = True

    doc.add_paragraph()

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("內帳系統完整架構 & 開發紀錄")
    run.font.size = Pt(16)

    for _ in range(3):
        doc.add_paragraph()

    info_lines = [
        "版本：v4.0",
        "日期：2026-04-11",
        "作者：陳柏諭",
        "狀態：機密文件",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)

    doc.add_page_break()

    # ===================================================================
    # 目錄
    # ===================================================================
    add_heading(doc, "目錄", level=1)
    toc_items = [
        "一、專案概要",
        "二、系統架構圖",
        "三、會計循環完整流程",
        "四、一鍵做賬 Pipeline（10 步驟）",
        "五、自動稽核模組",
        "六、財務分析模組",
        "七、排程器",
        "八、資料庫 Schema（22 張表）",
        "九、會計科目表（EAS）",
        "十、傳票分錄規則",
        "十一、稅務分類邏輯",
        "十二、GDrive 資料夾結構",
        "十三、OCR 三引擎",
        "十四、六宮格 Rich Menu 架構",
        "十五、CLI 工具",
        "十六、API 使用方式",
        "十七、測試",
        "十八、版本歷程",
        "附錄 A：開發意圖對照表",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ===================================================================
    # 一、專案概要
    # ===================================================================
    add_heading(doc, "一、專案概要", level=1)

    p = doc.add_paragraph()
    p.add_run("五家團膳公司的內帳系統，透過 LINE Bot 收據拍照 → OCR 辨識 → 記帳 → 歸檔 → 報表。").font.size = Pt(11)

    doc.add_paragraph()
    specs = [
        ("架構", "Python 3.10 + FastAPI + SQLite (WAL) + LINE Messaging API"),
        ("PM2 服務", "shanbot, port 8025"),
        ("外部 URL", "https://shanbot.kuangshin.tw"),
        ("資料庫", "data/shanbot.db（22 張表，複式簿記）"),
        ("GDrive", "/mnt/h/小魚資料/團膳公司資料/{公司}/"),
    ]
    for label, val in specs:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(val)
        run.font.size = Pt(11)

    doc.add_paragraph()
    add_heading(doc, "五家公司", level=2)
    companies = [
        ("1", "福利社", "升鼎商行", "81410187", "福利社/"),
        ("2", "王凱", "王凱食品有限公司", "90438334", "王凱/"),
        ("3", "台達2廠", "升鼎商行", "81410187", "台達2廠/"),
        ("4", "富燚", "富燚商行", "00281384", "富燚/"),
        ("5", "台達1廠", "升鼎商行", "81410187", "台達1廠/"),
    ]
    add_table_with_header(doc,
        ["ID", "簡稱", "全名", "統編", "GDrive 資料夾"],
        companies,
        col_widths=[1.5, 2.5, 5, 3, 3])

    doc.add_page_break()

    # ===================================================================
    # 二、系統架構圖
    # ===================================================================
    add_heading(doc, "二、系統架構圖", level=1)
    arch_code = """LINE Webhook -> main.py (FastAPI)
  |-- handlers/
  |   |-- postback_handler.py    29 路由
  |   |-- photo_handler.py       OCR 流程
  |   |-- file_handler.py        文件上傳
  |   +-- command_handler.py     文字指令
  |-- services/
  |   |-- accounting_service.py  會計循環核心
  |   |-- financial_report_service.py  四大報表
  |   |-- report_service.py      月報/年報
  |   |-- tax_export_service.py  稅務匯出
  |   |-- audit_service.py       自動稽核
  |   |-- pipeline_service.py    一鍵做賬
  |   |-- financial_analysis_service.py  財務分析
  |   |-- ocr_service.py         三引擎 OCR
  |   |-- flex_builder.py        Flex Message
  |   |-- line_service.py        LINE API
  |   |-- gdrive_service.py      GDrive 同步
  |   +-- market_service.py      農產品行情
  |-- state_manager.py           SQLite ORM (60+ CRUD)
  +-- task_manager.py            6 排程器"""
    add_code_block(doc, arch_code)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("核心流程：").bold = True
    flow_code = """照片上傳 (LINE)
  -> photo_handler.handle_photo_received()
  -> SHA256 去重
  -> OCR 三引擎（PaddleOCR -> Gemini VLM -> HunyuanOCR）
  -> 稅務自動分類（三聯/二聯/電子/免用 -> 扣抵代號）
  -> purchase_staging 暫存（status=pending）
  -> 使用者確認（好/修改/捨棄）
  -> confirm -> 傳票分錄（借：進貨+進項稅額 / 貸：現金）
  -> verify_balance() 借貸平衡檢查
  -> GDrive 歸檔（重新命名 YYMMDD_供應商_金額_#ID.jpg）
  -> 供應商子資料夾分類 + INDEX.csv 索引更新"""
    add_code_block(doc, flow_code)

    doc.add_page_break()

    # ===================================================================
    # 三、會計循環完整流程
    # ===================================================================
    add_heading(doc, "三、會計循環完整流程", level=1)

    p = doc.add_paragraph()
    p.add_run("遵循：台灣中小企業會計準則（EAS）+ 營業稅法 + 所得稅法").font.size = Pt(11)
    doc.add_paragraph()

    acct_steps = [
        ("1. 進貨分錄", "複式簿記核心。借：進貨 5110 + 進項稅額 1150 / 貸：現金 1100。根據稅務分類決定是否拆分稅額。"),
        ("2. 收入分錄", "借：現金 1100 / 貸：營業收入 4100 + 銷項稅額 2150。依公司每月營收登錄。"),
        ("3. 薪資分錄", "借：薪資 6110 + 勞健保 6200 + 勞退 6210 / 貸：現金 1100 + 代扣稅 2310。依員工薪資明細自動產生。"),
        ("4. 固定資產折舊", "直線法攤提。借：折舊費用 6160 / 貸：累計折舊 1500。每月自動計算。"),
        ("5. 期末結帳", "調整分錄 -> 結轉損益（4xxx/5xxx/6xxx 歸零到 3300 本期損益）-> 開立下期帳。"),
        ("6. 財務報表", "資產負債表、損益表、現金流量表、權益變動表（openpyxl 生成 Excel）。"),
        ("7. 營業稅摘要", "401 申報用。雙月彙總進項/銷項稅額，產出 MOF 格式 TXT 檔。"),
        ("8. Excel 帳冊", "8 Sheet：總帳、日記帳、試算表、損益表、資產負債表、科目餘額、明細帳、營業稅。"),
    ]
    for title, desc in acct_steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        run2 = p2.add_run(desc)
        run2.font.size = Pt(11)

    doc.add_page_break()

    # ===================================================================
    # 四、一鍵做賬 Pipeline
    # ===================================================================
    add_heading(doc, "四、一鍵做賬 Pipeline（10 步驟）", level=1)

    p = doc.add_paragraph()
    p.add_run("模組：services/pipeline_service.py").font.size = Pt(10)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("支援模式：單月/多月、全公司/單公司、自動確認/手動確認").font.size = Pt(11)
    doc.add_paragraph()

    pipeline_steps = [
        ("Step 1", "批次確認 + 生成分錄", "確認所有 pending 單據，生成進貨分錄 + 收入分錄（複式簿記）"),
        ("Step 2", "月報表", "採購彙總報表 + 憑證目錄（Excel）"),
        ("Step 3", "採購報告", "供應商統計、分類統計、每日採購金額"),
        ("Step 4", "會計帳冊", "8-sheet Excel（總帳/日記帳/試算表/損益表/資產負債表/科目餘額/明細帳/營業稅）"),
        ("Step 5", "四大財務報表", "資產負債表、損益表、現金流量表、權益變動表"),
        ("Step 6", "自動稽核", "6 項檢查（詳見第五章），結帳前確保數據完整性"),
        ("Step 7", "期末結帳", "調整分錄 → 結轉損益 → 月度會計總表更新（可選）"),
        ("Step 8", "稅務匯出", "MOF TXT + 會計 Excel + 經手人 PDF（雙月執行）"),
        ("Step 9", "財務分析報告", "成本結構 + 趨勢 + KPI + 風險 + 建議（Excel + 文字摘要）"),
        ("Step 10", "GDrive 歸檔", "所有產出檔案上傳至對應公司/月份資料夾"),
    ]
    add_table_with_header(doc,
        ["步驟", "名稱", "說明"],
        pipeline_steps,
        col_widths=[2, 4, 10])

    doc.add_paragraph()
    add_heading(doc, "Pipeline 回傳結構", level=2)
    add_code_block(doc, """{
    "year_month": "2026-03",
    "timestamp": "2026-03-31 20:15:00",
    "steps": [...],           # 每步驟結果
    "files_generated": [...], # 所有產出檔案路徑
    "audit_result": {...},    # 稽核摘要
    "overall_success": true   # 全部成功 or 部分異常
}""")

    doc.add_page_break()

    # ===================================================================
    # 五、自動稽核模組
    # ===================================================================
    add_heading(doc, "五、自動稽核模組", level=1)

    p = doc.add_paragraph()
    p.add_run("模組：services/audit_service.py").font.size = Pt(10)
    doc.add_paragraph()

    audit_items = [
        ("1. 試算表借貸平衡", "verify_trial_balance()", "驗證月度試算表 total_debit == total_credit（容許 $1 四捨五入差）"),
        ("2. 資產負債表恆等式", "verify_balance_sheet_equation()", "資產 = 負債 + 權益。結帳前加計未結轉淨利（4xxx - 5xxx - 6xxx）"),
        ("3. 進項稅額覈對", "verify_input_tax()", "purchase_staging 的 tax_amount 合計 vs journal_entries 進項稅額科目合計"),
        ("4. 每筆交易借貸平衡", "verify_entry_balance()", "逐筆檢查每一張傳票的借方 = 貸方"),
        ("5. 損益表交叉驗證", "verify_income_statement()", "收入 - 成本 - 費用 = 稅前淨利，與 monthly_accounting 對照"),
        ("6. 異常偵測", "detect_anomalies()", "高金額（>$50,000）、供應商集中度>50%、重複金額、未分類項目"),
    ]
    add_table_with_header(doc,
        ["檢查項目", "函數", "邏輯"],
        audit_items,
        col_widths=[4, 5, 8])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("稽核通過條件：").bold = True
    doc.add_paragraph("所有 balanced == True 且無 high-level anomaly", style='List Bullet')
    doc.add_paragraph("Pipeline 中 Step 6 稽核通過後才執行 Step 7 期末結帳", style='List Bullet')

    doc.add_page_break()

    # ===================================================================
    # 六、財務分析模組
    # ===================================================================
    add_heading(doc, "六、財務分析模組", level=1)

    p = doc.add_paragraph()
    p.add_run("模組：services/financial_analysis_service.py").font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "分析內容", level=2)
    analysis_items = [
        "成本結構分析：食材分類佔比排序 + 前 N 大供應商排名 + 每日採購趨勢",
        "環比趨勢分析：本月 vs 上月的採購成本、營收變動百分比",
        "經營指標（KPI）：毛利率、成本率、日均採購成本、可扣抵發票比例",
        "風險偵測（6 種）",
        "自動改進建議：根據風險類型生成對應可執行建議",
    ]
    for item in analysis_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_heading(doc, "6 種風險偵測", level=2)
    risk_items = [
        ("low_margin", "毛利率過低", "< 20% = high, < 30% = medium（行業標準 30-45%）"),
        ("cost_surge", "採購成本暴漲", "環比上升 > 20%"),
        ("supplier_concentration", "供應商集中度", ">= 50% = high, >= 35% = medium"),
        ("low_deductible", "可扣抵比例偏低", "< 60% 影響營業稅抵扣"),
        ("category_imbalance", "食材結構失衡", "單一分類 >= 60%"),
        ("revenue_decline", "收入下滑", "環比下降 > 10%"),
    ]
    add_table_with_header(doc,
        ["風險類型", "名稱", "觸發條件"],
        risk_items,
        col_widths=[4, 4, 8])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("產出格式：").bold = True
    p.add_run(" Excel 報告（含圖表）+ 文字摘要（可直接推送 LINE）").font.size = Pt(11)

    doc.add_page_break()

    # ===================================================================
    # 七、排程器
    # ===================================================================
    add_heading(doc, "七、排程器", level=1)

    p = doc.add_paragraph()
    p.add_run("模組：task_manager.py（asyncio 排程）").font.size = Pt(10)
    doc.add_paragraph()

    schedulers = [
        ("1", "HeartbeatScheduler", "每日 16:00", "推送當日採購摘要 + 待確認筆數到 LINE 群組"),
        ("2", "MarketSyncScheduler", "每日 07:00", "同步農產品批發市場行情（台北果菜市場 API）"),
        ("3", "MonthlySummaryScheduler", "每月 1 號 09:00", "上月採購總結報告 + 推送 LINE"),
        ("4", "MonthEndAnalysisScheduler", "每月最後一天 20:00", "自動偵測月底 → 完整做賬 Pipeline + 財務分析報告 + 推送 LINE（v4.0 新增）"),
        ("5", "WebhookGuardScheduler", "每 6 小時", "檢查 LINE webhook 連線狀態，異常自動重連"),
        ("6", "ExternalAPIGuardScheduler", "每 12 小時", "檢查外部 API（OCR / 市場行情）可用性"),
    ]
    add_table_with_header(doc,
        ["#", "排程器", "頻率", "功能"],
        schedulers,
        col_widths=[1, 5, 4, 7])

    doc.add_paragraph()
    add_heading(doc, "MonthEndAnalysisScheduler 執行流程", level=2)
    add_code_block(doc, """月底 20:00 觸發
  -> run_full_pipeline(year_month, auto_confirm=True, skip_tax_export=True)
  -> generate_monthly_analysis(year_month)
  -> generate_analysis_excel(year_month)
  -> LINE push: 財務分析摘要 + Pipeline 狀態
  -> 下次排程：下月最後一天 20:00""")

    doc.add_page_break()

    # ===================================================================
    # 八、資料庫 Schema
    # ===================================================================
    add_heading(doc, "八、資料庫 Schema（22 張表）", level=1)

    p = doc.add_paragraph()
    p.add_run("資料庫：SQLite (WAL mode) — data/shanbot.db").font.size = Pt(10)
    doc.add_paragraph()

    db_tables = [
        ("companies", "公司主檔", "id, short_name, full_name, tax_id, gdrive_folder"),
        ("ingredients", "食材主檔", "id, code, name, category, unit, current_price, market_ref_price"),
        ("suppliers", "供應商主檔", "id, name, tax_id, has_uniform_invoice, score"),
        ("purchase_staging", "採購暫存（OCR）", "id, company_id, supplier_name, invoice_*, total_amount, status, year_month"),
        ("purchase_items", "採購明細", "id, staging_id, item_name, quantity, unit_price, amount, category"),
        ("price_history", "價格歷史", "id, ingredient_id, price_date, source, avg_price"),
        ("recipes", "配方表", "id, name, category, servings, ingredient_cost"),
        ("recipe_ingredients", "配方明細（BOM）", "id, recipe_id, ingredient_id, quantity, unit"),
        ("menu_schedule", "菜單排程", "id, schedule_date, meal_type, slot, recipe_id"),
        ("monthly_cost", "月度成本結構", "id, company_id, year_month, ingredient_total, labor_total, overhead_total"),
        ("config", "系統配置", "key, value"),
        ("tax_exports", "稅務匯出記錄", "id, company_id, tax_period, export_type, file_path, total_amount"),
        ("account_mapping", "會計科目對照", "id, category, account_code, account_name"),
        ("income", "收入記錄", "id, company_id, year_month, amount, description"),
        ("conversation_state", "對話狀態", "chat_id, state, state_data"),
        ("financial_documents", "財務文件索引", "id, company_id, filename, doc_category, gdrive_path, status"),
        ("employees", "員工主檔", "id, company_id, name, position, base_salary, status"),
        ("payroll", "薪資明細", "id, employee_id, year_month, gross_salary, net_salary"),
        ("journal_entries", "分錄（日記帳）", "id, company_id, entry_date, source_type, account_code, debit, credit"),
        ("monthly_accounting", "月度會計總表", "id, company_id, year_month, total_income, net_profit, is_closed"),
        ("chart_of_accounts", "會計科目表", "code, name, category, normal_side"),
        ("fixed_assets", "固定資產", "id, company_id, name, cost, useful_life_months, depreciation_method"),
        ("report_confirmations", "報表確認追蹤", "id, company_id, period, report_type, status"),
    ]
    add_table_with_header(doc,
        ["表名", "用途", "關鍵欄位"],
        [(t[0], t[1], t[2]) for t in db_tables],
        col_widths=[4, 4, 9])

    doc.add_page_break()

    # ===================================================================
    # 九、會計科目表
    # ===================================================================
    add_heading(doc, "九、會計科目表（EAS 中小企業會計準則）", level=1)

    doc.add_paragraph()
    eas_accounts = [
        ("1xxx", "資產", "1100 現金及約當現金\n1150 進項稅額\n1200 應收帳款\n1300 存貨—原料\n1500 固定資產（含累計折舊）"),
        ("2xxx", "負債", "2100 應付帳款\n2150 銷項稅額\n2200 應付薪資\n2310 代扣所得稅"),
        ("3xxx", "權益", "3100 資本額\n3200 累積盈虧\n3300 本期損益"),
        ("4xxx", "收入", "4100 營業收入"),
        ("5xxx", "成本", "5110 進貨（食材）\n5120 直接人工\n5130 製造費用"),
        ("6xxx", "費用", "6110 薪資費用\n6120 租金支出\n6130 水電瓦斯費\n6140 保險費\n6150 運費\n6160 折舊費用\n6190 其他營業費用\n6200 勞健保費用\n6210 勞退費用"),
    ]
    add_table_with_header(doc,
        ["代碼範圍", "大類", "明細科目"],
        eas_accounts,
        col_widths=[3, 3, 11])

    doc.add_page_break()

    # ===================================================================
    # 十、傳票分錄規則
    # ===================================================================
    add_heading(doc, "十、傳票分錄規則", level=1)

    doc.add_paragraph()
    journal_rules = [
        ("進貨（可扣抵）", "進貨 5110 + 進項稅額 1150", "現金 1100"),
        ("進貨（不可扣抵）", "進貨 5110（含稅全額）", "現金 1100"),
        ("營業收入", "現金 1100", "營收 4100 + 銷項稅額 2150"),
        ("薪資", "薪資 6110 + 勞健保 6200 + 勞退 6210", "現金 1100 + 代扣稅 2310"),
        ("折舊", "折舊費用 6160", "累計折舊 1500"),
        ("期末結轉", "收入 4100 + 成本 5110 + 費用 6xxx", "本期損益 3300"),
    ]
    add_table_with_header(doc,
        ["交易類型", "借方", "貸方"],
        journal_rules,
        col_widths=[4, 6, 6])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("規則：").bold = True
    rules_list = [
        "每筆傳票必須 sum(debit) == sum(credit)，容許 $1 差額",
        "進貨可扣抵由稅務分類邏輯自動判斷",
        "薪資分錄根據 payroll 表自動產生",
        "期末結轉在 Pipeline Step 7 執行",
    ]
    for r in rules_list:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # ===================================================================
    # 十一、稅務分類邏輯
    # ===================================================================
    add_heading(doc, "十一、稅務分類邏輯", level=1)

    doc.add_paragraph()
    tax_rules = [
        ("有統編 + 三聯式", "21", "1", "可扣抵進項稅額"),
        ("有統編 + 電子發票", "25", "1", "可扣抵進項稅額"),
        ("有統編 + 二聯式", "22", "2", "不可扣抵（稅額計入進貨成本）"),
        ("無統編 / 免用發票", "22", "2", "不可扣抵（全額計入進貨成本）"),
    ]
    add_table_with_header(doc,
        ["條件", "格式代號", "扣抵代號", "說明"],
        tax_rules,
        col_widths=[5, 3, 3, 6])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("自動判定流程：").bold = True
    add_code_block(doc, """OCR 辨識發票
  -> 抓取統編（supplier_tax_id）
  -> 判斷發票類型（三聯/二聯/電子/免用）
  -> 自動設定 invoice_format_code + deduction_code
  -> 可扣抵：稅額獨立記帳到 1150 進項稅額
  -> 不可扣抵：稅額併入 5110 進貨成本""")

    doc.add_page_break()

    # ===================================================================
    # 十二、GDrive 資料夾結構
    # ===================================================================
    add_heading(doc, "十二、GDrive 資料夾結構", level=1)

    doc.add_paragraph()
    gdrive_tree = """/mnt/h/小魚資料/團膳公司資料/
+-- {公司}/                       (福利社/王凱/台達2廠/富燚/台達1廠)
    +-- 2026/
        +-- 03月/
        |   +-- 收據憑證/          拍照記帳的收據照片
        |   |   +-- {供應商}/      供應商子資料夾
        |   |   +-- INDEX.csv      憑證索引
        |   +-- 待確認/            未確認暫存
        |   +-- 採購單據/          expenditure 類文件
        |   +-- 月報表/            月報表匯出
        |   +-- 稅務匯出/          MOF TXT + 會計匯出
        |   +-- 會計資料/          revenue/financing/investment 等文件
        |   +-- 財務報表/          四大報表 Excel
        |   +-- 薪資表/            payroll 類文件
        |   +-- 菜單企劃/          菜單相關
        |   +-- 租約與合約/        general 類文件
        |   +-- INDEX_總覽.csv     月度總索引
        +-- 年度報表/
            +-- 食材價格對照/"""
    add_code_block(doc, gdrive_tree)

    doc.add_page_break()

    # ===================================================================
    # 十三、OCR 三引擎
    # ===================================================================
    add_heading(doc, "十三、OCR 三引擎", level=1)

    p = doc.add_paragraph()
    p.add_run("模組：services/ocr_service.py").font.size = Pt(10)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("三層 Fallback 機制，確保辨識率最大化：").font.size = Pt(11)
    doc.add_paragraph()

    ocr_engines = [
        ("1", "PaddleOCR", "本地", "免費、速度快、隱私保護", "手寫/模糊收據辨識率較低"),
        ("2", "Gemini VLM", "Google API", "多模態理解力強、結構化輸出", "需網路、有 API 額度"),
        ("3", "HunyuanOCR", "騰訊 API", "中文手寫辨識強", "備援引擎、延遲較高"),
    ]
    add_table_with_header(doc,
        ["優先序", "引擎", "部署方式", "優勢", "限制"],
        ocr_engines,
        col_widths=[2, 3, 3, 5, 5])

    doc.add_paragraph()
    add_heading(doc, "Fallback 邏輯", level=2)
    add_code_block(doc, """process_image(image_path):
  result = paddleocr_recognize(image_path)
  if result.confidence >= 60%:
      return result    # PaddleOCR 成功

  result = gemini_vlm_recognize(image_path)
  if result.confidence >= 60%:
      return result    # Gemini VLM 成功

  result = hunyuan_ocr_recognize(image_path)
  return result        # HunyuanOCR（最後手段）

信心度分級：
  >= 80%  -> AUTO_PASS（自動確認）
  60-79%  -> REVIEW（需人工審閱）
  < 60%   -> REJECT（建議重拍）""")

    doc.add_page_break()

    # ===================================================================
    # 十四、六宮格 Rich Menu 架構
    # ===================================================================
    add_heading(doc, "十四、六宮格 Rich Menu 架構", level=1)

    doc.add_paragraph()
    grid_layout = """+----------------+----------------+----------------+
|    格子 1       |    格子 2       |    格子 3       |
|  拍照記帳       |  財務資料       |  採購管理       |
|  menu=camera   |  menu=finance  |  menu=purchase |
|                |  _upload       |                |
+----------------+----------------+----------------+
|    格子 4       |    格子 5       |    格子 6       |
|  菜單企劃       |  報表生成       |  使用說明       |
|  menu=menu     |  menu=reports  |  menu=guide    |
|  _plan         |                |                |
+----------------+----------------+----------------+"""
    add_code_block(doc, grid_layout)

    doc.add_paragraph()
    grid_routes = [
        ("格子 1 拍照記帳", "menu=camera", "3 卡 Carousel（拍照訣竅 + 辨識流程 + 開始拍照）"),
        ("格子 2 財務資料", "menu=finance_upload", "4 卡 Carousel（說明 + 上傳 + 已上傳 + 確認統計）"),
        ("格子 3 採購管理", "menu=purchase", "1 卡 Bubble（待處理 + 市場行情 + 供應商 + 價格對照）"),
        ("格子 4 菜單企劃", "menu=menu_plan", "4 卡 Carousel（說明 + 菜單確認 + 圖片生成 + 成本試算）"),
        ("格子 5 報表生成", "menu=reports", "3 卡 Carousel（說明 + 四大報表 + 匯出功能）"),
        ("格子 6 使用說明", "menu=guide", "6 卡 Carousel（快速開始 + 步驟 + 群組指南 + 功能 + FAQ + 關於）"),
    ]
    add_table_with_header(doc,
        ["格子", "Postback Data", "Flex 結構"],
        grid_routes,
        col_widths=[4, 5, 8])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Postback Handler 路由數：29 個").bold = True

    doc.add_page_break()

    # ===================================================================
    # 十五、CLI 工具
    # ===================================================================
    add_heading(doc, "十五、CLI 工具", level=1)

    doc.add_paragraph()
    cli_tools = [
        ("批次 OCR + 重新命名", "python3 tools/batch_rename_receipts.py [--dry-run] [--limit N]", "掃描收據資料夾，OCR 辨識後重新命名為 YYMMDD_供應商_金額_#ID.jpg"),
        ("收據分類整理", "bash tools/organize_receipts.sh \"path\"", "依供應商名稱分類到子資料夾 + 更新 INDEX.csv"),
        ("Rich Menu 部署", "python3 scripts/deploy_richmenu.py", "上傳 Rich Menu 圖片 + 設定 postback action + 綁定用戶"),
        ("資料庫直查", "sqlite3 data/shanbot.db \"SQL\"", "查看待確認/借貸平衡/各公司數量/科目表"),
    ]
    add_table_with_header(doc,
        ["功能", "指令", "說明"],
        cli_tools,
        col_widths=[4, 6, 7])

    doc.add_page_break()

    # ===================================================================
    # 十六、API 使用方式
    # ===================================================================
    add_heading(doc, "十六、API 使用方式", level=1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("所有模組均可獨立呼叫，以下為常用 API：").font.size = Pt(11)
    doc.add_paragraph()

    add_heading(doc, "Pipeline", level=2)
    add_code_block(doc, """from services.pipeline_service import run_full_pipeline
result = run_full_pipeline("2026-03", auto_confirm=True)
# result["overall_success"] -> bool
# result["files_generated"] -> list[str]""")

    doc.add_paragraph()
    add_heading(doc, "稽核", level=2)
    add_code_block(doc, """from services.audit_service import (
    verify_trial_balance, verify_balance_sheet_equation,
    verify_input_tax, detect_anomalies
)
trial = verify_trial_balance("2026-03")
# trial["balanced"] -> bool
# trial["total_debit"] / trial["total_credit"]""")

    doc.add_paragraph()
    add_heading(doc, "財務分析", level=2)
    add_code_block(doc, """from services.financial_analysis_service import generate_monthly_analysis
analysis = generate_monthly_analysis("2026-03")
# analysis["kpi"]["gross_margin_pct"]
# analysis["risks"] -> list[dict]
# analysis["recommendations"] -> list[str]
# analysis["summary_text"] -> str (LINE 推送用)""")

    doc.add_paragraph()
    add_heading(doc, "會計", level=2)
    add_code_block(doc, """from services.accounting_service import (
    generate_journal_entries, verify_balance,
    generate_income_journal_entries, period_closing
)
generate_journal_entries(staging_id)
result = verify_balance(staging_id)
# result["balanced"] -> True, result["debit"] == result["credit"]""")

    doc.add_paragraph()
    add_heading(doc, "報表", level=2)
    add_code_block(doc, """from services.financial_report_service import (
    generate_balance_sheet, generate_income_statement,
    generate_cash_flow, generate_equity_changes
)
generate_balance_sheet("2026-03", output_dir="/tmp")

from services.tax_export_service import export_mof_txt, export_accounting_excel
export_mof_txt("2026-01-02", output_dir="/tmp")  # 期間: YYYY-MM 雙月""")

    doc.add_page_break()

    # ===================================================================
    # 十七、測試
    # ===================================================================
    add_heading(doc, "十七、測試", level=1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("測試框架：pytest").font.size = Pt(11)
    doc.add_paragraph()

    test_results = [
        ("通過", "680"),
        ("已知失敗", "6（known failures，不影響正式環境）"),
        ("跳過", "1（環境相依）"),
    ]
    add_table_with_header(doc,
        ["狀態", "數量"],
        test_results,
        col_widths=[4, 12])

    doc.add_paragraph()
    add_heading(doc, "已驗證勾稽鏈", level=2)
    verifications = [
        "8 筆採購 + 2 筆收入全部入帳成功",
        "試算表 315,250 = 315,250（借貸平衡）",
        "資產 99,500 = 負債 9,761 + 權益 89,739（恆等式成立）",
        "進項稅額覈對：staging $2,500 = journal $2,500",
        "每筆傳票逐筆驗證借貸平衡",
    ]
    for v in verifications:
        doc.add_paragraph(v, style='List Bullet')

    doc.add_page_break()

    # ===================================================================
    # 十八、版本歷程
    # ===================================================================
    add_heading(doc, "十八、版本歷程", level=1)

    doc.add_paragraph()
    versions = [
        ("v1.0", "2026-02", "單公司 LINE Bot + OCR 拍照記帳 + 基礎報表"),
        ("v2.0", "2026-02", "五公司多租戶架構 + 供應商管理 + 市場行情"),
        ("v2.2", "2026-02-18", "四大財務報表 + 財務文件上傳智能分類 + Rich Menu v2（六宮格）+ 610 測試通過"),
        ("v3.0", "2026-03", "複式簿記完整會計循環 + 傳票分錄 + 期末結帳 + 稅務匯出（MOF TXT）+ 會計帳冊 8 Sheet"),
        ("v4.0", "2026-04-11", "自動稽核（6 項驗證）+ 一鍵做賬 Pipeline（10 步驟）+ 財務分析報告（成本/趨勢/風險/建議）+ 月底自動排程 + 680 測試通過"),
    ]
    add_table_with_header(doc,
        ["版本", "日期", "重點"],
        versions,
        col_widths=[2, 3, 12])

    doc.add_page_break()

    # ===================================================================
    # 附錄 A：開發意圖對照表
    # ===================================================================
    add_heading(doc, "附錄 A：開發意圖對照表", level=1)

    p = doc.add_paragraph()
    p.add_run("每個功能模組的商業意圖對照：").font.size = Pt(11)
    doc.add_paragraph()

    intent_table = [
        ("OCR 拍照記帳", "員工不需手動輸入，降低進入門檻", "ocr_service + photo_handler"),
        ("複式分錄", "符合台灣 EAS 會計準則，可與外部會計對接", "accounting_service"),
        ("自動稽核", "降低人為記帳錯誤，確保報表正確性", "audit_service"),
        ("財務分析", "老闆不用看懂報表也能知道經營問題在哪", "financial_analysis_service"),
        ("月底自動排程", "零人工月結，每月最後一天自動完成全部帳務", "task_manager + pipeline_service"),
        ("稅務匯出", "直接對接記帳士/國稅局 401 申報格式", "tax_export_service"),
        ("一鍵做賬", "10 步驟一次跑完，不需逐一手動操作", "pipeline_service"),
        ("GDrive 歸檔", "所有文件雲端備份，多設備存取", "gdrive_service"),
        ("六宮格 Rich Menu", "非技術用戶也能直覺操作所有功能", "flex_builder + postback_handler"),
        ("多租戶隔離", "五家公司共用一個系統但數據完全隔離", "state_manager（company_id）"),
    ]
    add_table_with_header(doc,
        ["功能", "開發意圖", "對應模組"],
        intent_table,
        col_widths=[4, 7, 5])

    # ===================================================================
    # 儲存
    # ===================================================================
    return doc


if __name__ == "__main__":
    doc = generate_report()

    # 儲存到兩個路徑
    paths = [
        "/home/simon/shanbot/docs/小膳BOT_技術開發報告_v4.0.docx",
        "/mnt/h/小魚資料/團膳公司資料/小膳BOT_技術開發報告_v4.0.docx",
    ]

    for path in paths:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        doc.save(path)
        print(f"已儲存：{path}")

    print("\n完成！")
